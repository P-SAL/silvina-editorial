"""
word_exporter.py
Export analysis results to Word (.docx) format.
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from typing import Dict, Any


class WordExporter:
    """Export analysis results to formatted Word document."""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecute: pip install python-docx")
    
    def export_to_word(self, analysis_results: Dict[str, Any], output_path: str) -> bool:
        """
        Export analysis to Word document.
        
        Args:
            analysis_results: Dictionary with all analysis data
            output_path: Path where to save the .docx file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()
            
            # ============================================================
            # HEADER
            # ============================================================
            title = doc.add_heading('INFORME DE ANÁLISIS EDITORIAL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Silvina Editorial Assistant v0.7', level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # DOCUMENT INFO
            # ============================================================
            doc.add_heading('📄 Información del Documento', level=1)
            info = analysis_results['document_info']
            
            doc.add_paragraph(f"Archivo: {analysis_results['filename']}")
            doc.add_paragraph(f"Título: {info.get('title', 'No especificado')}")
            
            if info.get('authors'):
                doc.add_paragraph(f"Autores: {info['authors']}")
                    
                                
            doc.add_paragraph(f"Total de palabras: {info['word_count']:,}")
            doc.add_paragraph(f"Páginas estimadas: {info['estimated_pages']}")
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # PUBLISHABILITY DECISION - MOST IMPORTANT SECTION
            # ============================================================
            doc.add_heading('🎯 DECISIÓN DE PUBLICACIÓN', level=1)
            
            can_publish, reason = self._determine_publishability(analysis_results)
            
            # Create decision paragraph with colored text
            decision_para = doc.add_paragraph()
            decision_run = decision_para.add_run(
                "✅ RECOMENDACIÓN: APTO PARA PUBLICACIÓN" if can_publish 
                else "❌ RECOMENDACIÓN: REQUIERE REVISIÓN ANTES DE PUBLICACIÓN"
            )
            decision_run.bold = True
            decision_run.font.size = Pt(14)
            decision_run.font.color.rgb = RGBColor(0, 128, 0) if can_publish else RGBColor(192, 0, 0)
            
            doc.add_paragraph()
            
            # Justification
            justification = doc.add_paragraph()
            justification.add_run("Justificación: ").bold = True
            justification.add_run(reason)
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # CLASSIFICATION
            # ============================================================
            doc.add_heading('🏷️ Clasificación del Artículo', level=1)
            classification = analysis_results['classification']
            
            category_name = self._format_category(classification['category'])
            doc.add_paragraph(f"Categoría: {category_name}")
            doc.add_paragraph(f"Confianza: {classification['confidence']:.1%}")
            
            if classification.get('reasoning'):
                reasoning = doc.add_paragraph()
                reasoning.add_run("Razonamiento: ").bold = True
                reasoning.add_run(classification['reasoning'])
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # QUALITY ANALYSIS
            # ============================================================
            doc.add_heading('⭐ Análisis de Calidad', level=1)
            quality = analysis_results['quality_analysis']
            
            # Overall score
            score_para = doc.add_paragraph()
            score_para.add_run("Puntuación general: ").bold = True
            score_run = score_para.add_run(f"{quality['overall_score']:.1f}/10.0")
            score_run.bold = True
            
            # Quality level
            level_text = self._format_quality_level(quality['quality_level'])
            level_para = doc.add_paragraph()
            level_para.add_run("Nivel de calidad: ").bold = True
            level_para.add_run(level_text)
            
            # Dimension scores
            if quality.get('dimensions'):
                doc.add_paragraph()
                doc.add_paragraph("Puntuación por dimensión:").bold = True
                
                for dim_name, dim_data in quality['dimensions'].items():
                    score = dim_data.get('score', 0)
                    feedback = dim_data.get('feedback', '')
                    
                    dim_para = doc.add_paragraph(style='List Bullet')
                    dim_para.add_run(f"{dim_name.capitalize()}: ").bold = True
                    dim_para.add_run(f"{score:.1f}/10.0")
                    
                    if feedback:
                        feedback_para = doc.add_paragraph(f"  → {feedback}", style='List Bullet 2')
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # STRUCTURE VALIDATION
            # ============================================================
            doc.add_heading('📋 Validación de Estructura', level=1)
            structure = analysis_results['structure_validation']
            
            # Status
            status_para = doc.add_paragraph()
            if structure['is_valid']:
                status_run = status_para.add_run("✓ Estructura válida según normas EUMIC")
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                status_run = status_para.add_run("✗ Estructura incompleta")
                status_run.font.color.rgb = RGBColor(192, 0, 0)
            status_run.bold = True
            
            # Missing sections
            if structure.get('missing_sections'):
                doc.add_paragraph()
                doc.add_paragraph("Secciones faltantes:").bold = True
                for section in structure['missing_sections']:
                    doc.add_paragraph(f"• {section}", style='List Bullet')
            
            # Section details
            if structure.get('details'):
                doc.add_paragraph()
                doc.add_paragraph("Detalle de secciones:").bold = True
                for section_name, section_info in structure['details'].items():
                    status_icon = "✓" if section_info.get('present') else "✗"
                    doc.add_paragraph(f"{status_icon} {section_name}", style='List Bullet')
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # CITATIONS ANALYSIS
            # ============================================================
            doc.add_heading('📚 Análisis de Citas y Referencias', level=1)
            citations = analysis_results['citations_analysis']
            
            doc.add_paragraph(f"Total de citas en el texto: {citations['total_citations']}")
            doc.add_paragraph(f"Total de referencias bibliográficas: {citations['total_references']}")
            doc.add_paragraph(f"Citas con referencia encontrada: {citations['matched_count']}")
            doc.add_paragraph(f"Citas sin referencia: {citations['unmatched_count']}")
            
            # Match rate
            if citations['total_citations'] > 0:
                match_rate = citations['matched_count'] / citations['total_citations'] * 100
                rate_para = doc.add_paragraph()
                rate_para.add_run("Tasa de coincidencia: ").bold = True
                rate_run = rate_para.add_run(f"{match_rate:.1f}%")
                
                if match_rate >= 90:
                    rate_run.font.color.rgb = RGBColor(0, 128, 0)
                elif match_rate >= 70:
                    rate_run.font.color.rgb = RGBColor(255, 140, 0)
                else:
                    rate_run.font.color.rgb = RGBColor(192, 0, 0)
            
            # Citations by type
            if citations.get('by_type'):
                doc.add_paragraph()
                doc.add_paragraph("Distribución por tipo de cita:").bold = True
                for cite_type, count in citations['by_type'].items():
                    doc.add_paragraph(f"• {cite_type}: {count}", style='List Bullet')
            
            # Unmatched citations (first 10)
            if citations.get('unmatched_citations'):
                doc.add_paragraph()
                doc.add_paragraph("Primeras citas sin referencia (máximo 10):").bold = True
                
                for i, unmatched in enumerate(citations['unmatched_citations'][:10], 1):
                    citation_text = unmatched
                    doc.add_paragraph(f"{i}. {citation_text}", style='List Bullet')
            
            doc.add_paragraph()  # Spacing
            
            # ============================================================
            # RECOMMENDATIONS
            # ============================================================
            if analysis_results.get('recommendations'):
                doc.add_heading('💡 Recomendaciones', level=1)
                
                recommendations = analysis_results['recommendations']
                
                # Group by priority
                high_priority = [r for r in recommendations if r.get('priority') == 'alta']
                medium_priority = [r for r in recommendations if r.get('priority') == 'media']
                low_priority = [r for r in recommendations if r.get('priority') == 'baja']
                
                if high_priority:
                    doc.add_paragraph("Prioridad Alta:").bold = True
                    for rec in high_priority:
                        para = doc.add_paragraph(rec['message'], style='List Bullet')
                        para.runs[0].font.color.rgb = RGBColor(192, 0, 0)
                
                if medium_priority:
                    doc.add_paragraph()
                    doc.add_paragraph("Prioridad Media:").bold = True
                    for rec in medium_priority:
                        para = doc.add_paragraph(rec['message'], style='List Bullet')
                        para.runs[0].font.color.rgb = RGBColor(255, 140, 0)
                
                if low_priority:
                    doc.add_paragraph()
                    doc.add_paragraph("Prioridad Baja:").bold = True
                    for rec in low_priority:
                        doc.add_paragraph(rec['message'], style='List Bullet')
            
            # ============================================================
            # FOOTER
            # ============================================================
            doc.add_paragraph()
            doc.add_paragraph("─" * 60)
            footer = doc.add_paragraph()
            footer.add_run("Generado por Silvina Editorial Assistant v0.7").italic = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"      ❌ Error al crear documento Word: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _determine_publishability(self, results: Dict[str, Any]) -> tuple:
        """
        Determine if document can be published based on analysis.
        
        Returns:
            (can_publish: bool, reason: str)
        """
        quality_score = results['quality_analysis']['overall_score']
        is_structure_valid = results['structure_validation']['is_valid']
        citations_matched = results['citations_analysis']['matched_count']
        citations_total = results['citations_analysis']['total_citations']
        
        # Calculate citation match rate
        if citations_total > 0:
            citation_rate = citations_matched / citations_total
        else:
            citation_rate = 1.0  # No citations is acceptable for some article types
        
        # Decision logic based on EUMIC standards
        if quality_score >= 7.0 and is_structure_valid and citation_rate >= 0.9:
            return (
                True, 
                f"El documento cumple con los estándares de calidad (puntuación: {quality_score:.1f}/10), "
                f"estructura completa, y referencias adecuadas ({citation_rate:.1%} de coincidencia) "
                "requeridos por las normas EUMIC y APA 7."
            )
        
        elif quality_score >= 6.5 and is_structure_valid and citation_rate >= 0.85:
            return (
                False, 
                f"Calidad aceptable ({quality_score:.1f}/10) y estructura válida, pero requiere "
                f"mejoras menores en las referencias (tasa de coincidencia: {citation_rate:.1%}). "
                "Se recomienda revisión antes de publicación."
            )
        
        elif not is_structure_valid:
            missing = ", ".join(results['structure_validation']['missing_sections'])
            return (
                False, 
                f"Estructura incompleta según normas EUMIC. Faltan las siguientes secciones obligatorias: {missing}. "
                f"Calidad actual: {quality_score:.1f}/10."
            )
        
        elif citation_rate < 0.8:
            return (
                False,
                f"Inconsistencias importantes en las referencias bibliográficas. "
                f"Solo {citation_rate:.1%} de las citas tienen referencia correspondiente. "
                f"Revise el formato APA 7 y asegure que todas las citas estén referenciadas."
            )
        
        else:
            return (
                False, 
                f"Calidad insuficiente ({quality_score:.1f}/10). El documento requiere revisión "
                "sustancial en claridad, coherencia, argumentación y/o metodología antes de publicación."
            )
    
    def _format_category(self, category: str) -> str:
        """Format category name for display."""
        category_names = {
            'RESEARCH_ARTICLE': 'Artículo Científico (Investigación)',
            'REVIEW_ARTICLE': 'Artículo de Revisión',
            'REFLECTION_ARTICLE': 'Artículo de Reflexión',
            'SHORT_ARTICLE': 'Artículo Corto',
            'CASE_REPORT': 'Reporte de Caso',
            'DIVULGATION': 'Artículo de Divulgación',
            'UNKNOWN': 'No Clasificado'
        }
        return category_names.get(category, category)
    
    def _format_quality_level(self, quality_level):
        """
        Convert QualityLevel enum to Spanish readable text.
        ALWAYS returns a string.
        """
        level_names = {
            'EXCELLENT': 'Excelente',
            'GOOD': 'Bueno',
            'ACCEPTABLE': 'Aceptable',
            'NEEDS_IMPROVEMENT': 'Necesita Mejoras',
            'POOR': 'Deficiente'
        }

        # Case 1: Enum → use .name
        if hasattr(quality_level, "name"):
            return level_names.get(quality_level.name, quality_level.name)

        # Case 2: Already a string
        if isinstance(quality_level, str):
            return level_names.get(quality_level, quality_level)

        # Fallback (safety)
        return str(quality_level)
