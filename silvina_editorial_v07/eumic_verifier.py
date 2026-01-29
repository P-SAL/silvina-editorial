"""
eumic_verifier.py
Verifies EUMIC editorial standards compliance
Part of Silvina Editorial Assistant v0.7
"""
from __future__ import annotations

from typing import List, Dict, Any
from dataclasses import dataclass
from enum import Enum
import re


class EumicSeverity(Enum):
    """Severity levels for EUMIC violations"""
    CRITICAL = "🔴 CRÍTICO"
    WARNING = "🟡 ADVERTENCIA"
    INFO = "ℹ️ INFO"


@dataclass
class EumicViolation:
    """Represents a EUMIC standard violation"""
    category: str
    message: str
    severity: EumicSeverity
    details: str = ""


class EumicVerifier:
    """Verifies document compliance with EUMIC editorial standards."""
    
    def __init__(self):
        self.violations: List[EumicViolation] = []
    
    def verify_document(self, doc, document_content) -> List[EumicViolation]:
        """
        Run all EUMIC verification checks.
        
        Args:
            doc: python-docx Document object
            document_content: DocumentContent object
            
        Returns:
            List of violations found (empty if all compliant)
        """
        self.violations = []
        
        # Run all verification checks
        self._verify_format(doc)
        self._verify_figures(doc)
        self._verify_tables(doc)
        self._verify_formulas(doc)
        self._verify_abstract_keywords(doc, document_content)
        
        return self.violations
    
    def _verify_format(self, doc):
        """1. Formato general del documento"""
        from docx.shared import Pt, Cm
        
        # Check margins (2.5 cm requirement)
        sections = doc.sections
        if sections:
            section = sections[0]
            margin_cm = 2.5
            tolerance_cm = 0.3  # Allow small deviation
            
            required_twips = Cm(margin_cm).twips
            tolerance_twips = Cm(tolerance_cm).twips
            
            margins_to_check = [
                ("superior", section.top_margin),
                ("inferior", section.bottom_margin),
                ("izquierdo", section.left_margin),
                ("derecho", section.right_margin)
            ]
            
            for margin_name, margin_value in margins_to_check:
                if abs(margin_value.twips - required_twips) > tolerance_twips:
                    actual_cm = margin_value.cm
                    self.violations.append(EumicViolation(
                        category="Formato General",
                        message=f"Margen {margin_name} no cumple estándar EUMIC",
                        severity=EumicSeverity.WARNING,
                        details=f"Requerido: 2.5 cm, Actual: {actual_cm:.2f} cm"
                    ))
        
        # Check font (Times New Roman or Arial, 12pt)
        valid_fonts = ["Times New Roman", "Arial", "Calibri"]
        
        fonts_used = set()
        sizes_used = set()
        
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                if run.font.size:
                    sizes_used.add(run.font.size)
        
        # Check if any non-standard fonts are used
        non_standard_fonts = fonts_used - set(valid_fonts)
        if non_standard_fonts:
            self.violations.append(EumicViolation(
                category="Formato General",
                message="Fuentes no estándar detectadas",
                severity=EumicSeverity.WARNING,
                details=f"Usar Times New Roman o Arial. Detectadas: {', '.join(non_standard_fonts)}"
            ))
        
        # Check font sizes (should be mostly 12pt)
        non_standard_sizes = [s for s in sizes_used if s and abs(s.pt - 12) > 1]
        if len(non_standard_sizes) > 0:
            sizes_str = ', '.join([f"{s.pt:.0f}pt" for s in non_standard_sizes[:3]])
            # This is often OK for titles/headers, so make it INFO level
            self.violations.append(EumicViolation(
                category="Formato General",
                message="Tamaños de fuente variables detectados",
                severity=EumicSeverity.INFO,
                details=f"Predominantemente use 12pt. Detectados: {sizes_str}"
            ))
        
        # Check text alignment (should be justified)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        non_justified = 0
        total_paras = 0
        
        for para in doc.paragraphs:
            if para.text.strip():  # Only check non-empty paragraphs
                total_paras += 1
                if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    non_justified += 1
        
        if total_paras > 0 and non_justified / total_paras > 0.3:  # More than 30% not justified
            self.violations.append(EumicViolation(
                category="Formato General",
                message="Texto no está completamente justificado",
                severity=EumicSeverity.WARNING,
                details=f"{non_justified}/{total_paras} párrafos no justificados"
            ))
    
    def _verify_figures(self, doc):
        """2. Figuras"""
        # Count images/figures
        image_count = 0
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
        except:
            pass  # Skip if no relationships
        
        if image_count == 0:
            return  # No figures to verify
        
        # Check for figure captions (look for "Figura" or "Fig." in paragraphs)
        figure_captions = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.lower().startswith(('figura', 'fig.', 'figure')):
                figure_captions.append(text)
        
        if len(figure_captions) < image_count:
            self.violations.append(EumicViolation(
                category="Figuras",
                message="Figuras sin título descriptivo",
                severity=EumicSeverity.WARNING,
                details=f"{image_count} imágenes detectadas, {len(figure_captions)} títulos encontrados"
            ))
        
        # Check numbering (Figura 1, Figura 2, etc.)
        numbered_correctly = True
        expected_num = 1
        
        for caption in figure_captions:
            match = re.search(r'figura\s+(\d+)', caption.lower())
            if match:
                num = int(match.group(1))
                if num != expected_num:
                    numbered_correctly = False
                    break
                expected_num += 1
        
        if not numbered_correctly and len(figure_captions) > 1:
            self.violations.append(EumicViolation(
                category="Figuras",
                message="Numeración de figuras inconsistente",
                severity=EumicSeverity.WARNING,
                details="Las figuras deben numerarse consecutivamente (Figura 1, Figura 2, ...)"
            ))
    
    def _verify_tables(self, doc):
        """3. Tablas"""
        tables = doc.tables
        
        if len(tables) == 0:
            return  # No tables to verify
        
        # Check for table titles (should appear BEFORE the table in Word)
        table_titles = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.lower().startswith(('tabla', 'table', 'cuadro')):
                table_titles.append(text)
        
        if len(table_titles) < len(tables):
            self.violations.append(EumicViolation(
                category="Tablas",
                message="Tablas sin título descriptivo",
                severity=EumicSeverity.WARNING,
                details=f"{len(tables)} tablas detectadas, {len(table_titles)} títulos encontrados. Los títulos deben estar en la parte superior."
            ))
        
        # Check numbering
        numbered_correctly = True
        expected_num = 1
        
        for title in table_titles:
            match = re.search(r'tabla\s+(\d+)', title.lower())
            if match:
                num = int(match.group(1))
                if num != expected_num:
                    numbered_correctly = False
                    break
                expected_num += 1
        
        if not numbered_correctly and len(table_titles) > 1:
            self.violations.append(EumicViolation(
                category="Tablas",
                message="Numeración de tablas inconsistente",
                severity=EumicSeverity.WARNING,
                details="Las tablas deben numerarse consecutivamente (Tabla 1, Tabla 2, ...)"
            ))
    
    def _verify_formulas(self, doc):
        """4. Fórmulas"""
        # Detect formulas (OMath objects in Word)
        formula_count = 0
        
        for para in doc.paragraphs:
            # Check for equation fields
            for run in para.runs:
                # FIXED: Check if run contains equation (string comparison, not bytes)
                try:
                    xml_str = run._element.xml
                    if isinstance(xml_str, bytes):
                        xml_str = xml_str.decode('utf-8')
                    if '<m:oMath' in xml_str or '<w:equation' in xml_str:
                        formula_count += 1
                        break
                except:
                    # If XML parsing fails, skip this run
                    continue
        
        if formula_count == 0:
            return  # No formulas to verify
        
        # Check alignment (formulas should typically be centered)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        unaligned_formulas = 0
        for para in doc.paragraphs:
            has_formula = False
            for run in para.runs:
                try:
                    xml_str = run._element.xml
                    if isinstance(xml_str, bytes):
                        xml_str = xml_str.decode('utf-8')
                    if '<m:oMath' in xml_str or '<w:equation' in xml_str:
                        has_formula = True
                        break
                except:
                    continue
            
            if has_formula and para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                unaligned_formulas += 1
        
        if unaligned_formulas > 0:
            self.violations.append(EumicViolation(
                category="Fórmulas",
                message="Fórmulas no centradas",
                severity=EumicSeverity.INFO,
                details=f"{unaligned_formulas}/{formula_count} fórmulas no están centradas"
            ))
    
    def _verify_abstract_keywords(self, doc, document_content):
        """5. Resumen y palabras clave"""
        
        # Check for abstract/resumen
        has_abstract = False
        abstract_word_count = 0
        
        # FIXED: Use enumerate instead of index() to avoid ValueError
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if any(keyword in text.lower() for keyword in ['resumen', 'abstract', 'síntesis']):
                has_abstract = True
                # Try to count words in next few paragraphs (likely the abstract content)
                abstract_text = ""
                for i in range(para_idx, min(para_idx + 5, len(doc.paragraphs))):
                    abstract_text += " " + doc.paragraphs[i].text
                abstract_word_count = len(abstract_text.split())
                break
        
        if not has_abstract:
            self.violations.append(EumicViolation(
                category="Resumen y Palabras Clave",
                message="Falta sección de Resumen/Abstract",
                severity=EumicSeverity.CRITICAL,
                details="El documento debe incluir un resumen de 150-250 palabras"
            ))
        elif abstract_word_count < 100 or abstract_word_count > 300:
            self.violations.append(EumicViolation(
                category="Resumen y Palabras Clave",
                message="Extensión del resumen fuera de rango",
                severity=EumicSeverity.WARNING,
                details=f"Requerido: 150-250 palabras. Detectado: ~{abstract_word_count} palabras"
            ))
        
        # Check for keywords
        has_keywords = False
        keyword_count = 0
        
        for para in doc.paragraphs:
            text = para.text.lower()
            if any(kw in text for kw in ['palabras clave', 'keywords', 'key words']):
                has_keywords = True
                # Count keywords (separated by commas or semicolons)
                keyword_text = text.split(':', 1)[-1] if ':' in text else text
                keyword_count = len([k for k in re.split(r'[,;]', keyword_text) if k.strip()])
                break
        
        if not has_keywords:
            self.violations.append(EumicViolation(
                category="Resumen y Palabras Clave",
                message="Faltan palabras clave",
                severity=EumicSeverity.CRITICAL,
                details="Se requieren 3-5 palabras clave relevantes al contenido"
            ))
        elif keyword_count < 3 or keyword_count > 5:
            self.violations.append(EumicViolation(
                category="Resumen y Palabras Clave",
                message="Número incorrecto de palabras clave",
                severity=EumicSeverity.WARNING,
                details=f"Requerido: 3-5 palabras clave. Detectado: {keyword_count}"
            ))
    
    def format_violations_report(self, violations: List[EumicViolation]) -> str:
        """
        Format violations into readable report.
        Only returns content if there are violations.
        """
        if not violations:
            return ""
        
        report = "\n📋 VERIFICACIÓN EUMIC:\n"
        report += "=" * 80 + "\n"
        
        # Group by severity
        critical = [v for v in violations if v.severity == EumicSeverity.CRITICAL]
        warnings = [v for v in violations if v.severity == EumicSeverity.WARNING]
        info = [v for v in violations if v.severity == EumicSeverity.INFO]
        
        # Critical violations
        if critical:
            report += f"\n🔴 CRÍTICO ({len(critical)}):\n"
            for i, v in enumerate(critical, 1):
                report += f"   {i}. [{v.category}] {v.message}\n"
                if v.details:
                    report += f"      → {v.details}\n"
        
        # Warning violations
        if warnings:
            report += f"\n🟡 ADVERTENCIAS ({len(warnings)}):\n"
            for i, v in enumerate(warnings, 1):
                report += f"   {i}. [{v.category}] {v.message}\n"
                if v.details:
                    report += f"      → {v.details}\n"
        
        # Info violations
        if info:
            report += f"\nℹ️  INFORMACIÓN ({len(info)}):\n"
            for i, v in enumerate(info, 1):
                report += f"   {i}. [{v.category}] {v.message}\n"
                if v.details:
                    report += f"      → {v.details}\n"
        
        report += "=" * 80 + "\n"
        
        return report


def verify_eumic_compliance(doc, document_content) -> str:
    """
    Convenience function to verify EUMIC compliance.
    Returns formatted report string (empty if compliant).
    """
    verifier = EumicVerifier()
    violations = verifier.verify_document(doc, document_content)
    return verifier.format_violations_report(violations)
