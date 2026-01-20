"""
main.py
Main entry point for Silvina Editorial Assistant v0.7
Orchestrates the complete document analysis workflow.
"""

import sys
import os
from pathlib import Path
from typing import Optional, Dict, Any
import json
from docx import Document
from domain.enums import ArticleType, QualityLevel

# Add project root to path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

# Import configuration
from presentation.config import Config

# Import domain models
from domain.enums import ClassificationCategory, QualityLevel
from domain.enums import ArticleType
from domain.enums import ArticleSize


# Import data access layer
from data_access.word_reader import WordReader
from data_access.content_extractor import ContentExtractor
from data_access.citation_parser import CitationParser
from data_access.reference_parser import ReferenceParser

# Import business logic layer
from business_logic.article_classifier import ArticleClassifier
from business_logic.quality_analyzer import QualityAnalyzer
from business_logic.citation_matcher import CitationMatcher
from business_logic.structure_validator import StructureValidator

# Import presentation layer
from presentation.report_formatter import ReportFormatter
from presentation.word_exporter import WordExporter, DOCX_AVAILABLE


class SilvinaEditorialAssistant:
    """Main orchestrator for the Silvina Editorial Assistant."""
    
    def __init__(self, config: Optional[Config] = None):
        """
        Initialize Silvina with configuration.
        
        Args:
            config: Configuration object. If None, loads default config.
        """
        self.config = config or Config()
        
        print("🔧 Inicializando Silvina Editorial Assistant v0.7...")
        
        # Initialize components
        try:
            self.word_reader = WordReader()
            self.content_extractor = ContentExtractor()
            self.citation_parser = CitationParser()
            self.reference_parser = ReferenceParser()
            
            self.article_classifier = ArticleClassifier(
                model_name=self.config.ollama_model,
                base_url=self.config.ollama_base_url
            )
            self.quality_analyzer = QualityAnalyzer(
                model_name=self.config.ollama_model,
                base_url=self.config.ollama_base_url
            )
            self.citation_matcher = None  # Will be initialized when we have data
            self.structure_validator = StructureValidator()
            
            self.report_formatter = ReportFormatter()
            self.word_exporter = WordExporter() if DOCX_AVAILABLE else None
            
            print("✅ Silvina inicializada correctamente\n")
            
        except Exception as e:
            print(f"❌ Error al inicializar Silvina: {e}")
            raise
    
    def analyze_document(self, document_path: str) -> Dict[str, Any]:
        """
        Perform complete analysis of a document.
        
        Args:
            document_path: Path to the Word document (.docx)
            
        Returns:
            Dictionary containing all analysis results
        """
        print(f"📄 Analizando documento: {Path(document_path).name}")
        print("=" * 80)
        
        try:
            # Step 1: Read document
            print("\n[1/7] 📖 Leyendo documento...")
            doc = Document(document_path)
            document_text = "\n".join(p.text for p in doc.paragraphs)
            paragraphs = self.word_reader.read_word_document(document_path)
            if not paragraphs:
                raise ValueError("El documento está vacío o no se pudo leer")
            print(f"      ✓ {len(paragraphs)} párrafos leídos")
            
            # Step 2: Extract content
            print("\n[2/7] 🔍 Extrayendo contenido estructurado...")
            document_content = self.content_extractor.extract_content(paragraphs)
            print(f"      ✓ Título: {document_content.title or 'No especificado'}")
            print(f"      ✓ Total de palabras: {document_content.word_count:,}")
            
            # Step 3: Parse citations and references
            print("\n[3/7] 📚 Analizando citas y referencias...")
            # Extract citations from all paragraphs
            citations = []
            for idx, paragraph in enumerate(paragraphs):
                citations.extend(self.citation_parser.parse(paragraph, idx))
            # Extract references from bibliography section
            bibliography_text, section_type = self.reference_parser.extract_from_paragraphs(paragraphs)
            references, _ = self.reference_parser.parse_section(bibliography_text)
            print(f"      ✓ {len(citations)} citas encontradas")
            print(f"      ✓ {len(references)} referencias encontradas")
            
            # Step 4: Classify article
            print("\n[4/7] 🏷️  Clasificando tipo de artículo...")
            classification = self.article_classifier.classify_article(document_content)
            
            category_name = self._format_category(classification.article_type)
            print(f"      ✓ Categoría: {category_name}")
            print(f"      ✓ Confianza: {classification.confidence:.1%}")
            
            # Step 5: Analyze quality
            print("\n[5/7] ⭐ Analizando calidad...")
            quality_result = self.quality_analyzer.analyze_quality(
                document_content,
                classification
            )
            print(f"      ✓ Puntuación: {quality_result.overall_score:.1f}/10.0")
            print(f"      ✓ Nivel: {self._format_quality_level(quality_result.quality_level)}")
            
            # Step 6: Validate structure
            print("\n[6/7] 📋 Validando estructura...")
            structure_result = self.structure_validator.validate_structure(
                document_content,
                classification.article_type
            )
            status = "✓ VÁLIDA" if structure_result.is_valid else "✗ INCOMPLETA"
            print(f"      {status}")
            if structure_result.missing_sections:
                print(f"      ⚠ Secciones faltantes: {len(structure_result.missing_sections)}")
            
            # Step 7: Match citations
            print("\n[7/7] 🔗 Relacionando citas con referencias...")
            # Initialize citation matcher with extracted data
            self.citation_matcher = CitationMatcher(citations, references)
            citation_analysis = self.citation_matcher.match_citations_to_references(
                section_type
            )

            match_rate = (citation_analysis.matched_count / citation_analysis.total_citations * 100 
                         if citation_analysis.total_citations > 0 else 0)
            print(f"      ✓ Tasa de coincidencia: {match_rate:.1f}%")
            
            # Compile results
            print("\n" + "=" * 80)
            print("✅ Análisis completado exitosamente\n")
            
            analysis_results = {
                'filename': Path(document_path).name,
                'document_info': {
                    'title': document_content.title,
                    'authors': document_content.authors,
                    'word_count': document_content.word_count,
                    'paragraph_count': len(document_content.paragraphs),
                    'estimated_pages': document_content.word_count // 250  # ~250 words per page
                },
                'classification': {
                    'category': classification.article_type,
                    'article_size': classification.article_size,
                    'confidence': classification.confidence,
                    'reasoning': classification.reasoning
                },
     
                'quality_analysis': {
                    'overall_score': quality_result.overall_score,
                    'quality_level': quality_result.quality_level,
                    'dimensions': quality_result.dimension_scores
                },
                'structure_validation': {
                    'is_valid': structure_result.is_valid,
                    'missing_sections': structure_result.missing_sections,
                    'details': structure_result.section_details
                },
                'citations_analysis': {
                    'total_citations': citation_analysis.total_citations,
                    'total_references': citation_analysis.total_references,
                    'matched_count': citation_analysis.matched_count,
                    'unmatched_count': citation_analysis.unmatched_count,
                    'by_type': citation_analysis.citations_by_type,
                    'unmatched_citations': [
                        {'citation_text': c} if isinstance(c, str) else c
                        for c in citation_analysis.unmatched_citations[:20]
                    ]
                   
                },
                'recommendations': self._generate_recommendations(
                    classification,
                    quality_result,
                    structure_result,
                    citation_analysis
                )
            }
            
            return analysis_results
            
        except Exception as e:
            print(f"\n❌ Error durante el análisis: {e}")
            raise
    
    def save_text_report(self, analysis_results: Dict[str, Any], output_path: str):
        """Save analysis results as a text report."""
        try:
            print(f"💾 Guardando reporte de texto: {output_path}")
            report_text = self.report_formatter.generate_full_report(analysis_results)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(report_text)
            
            print(f"   ✅ Reporte guardado exitosamente")
            
        except Exception as e:
            print(f"   ❌ Error al guardar reporte: {e}")
            raise
    
    def save_word_report(self, analysis_results: Dict[str, Any], output_path: str):
        """Save analysis results as a Word document."""
        if not self.word_exporter:
            print("   ⚠ Word export no disponible (falta instalar python-docx)")
            return False
        
        try:
            print(f"💾 Guardando reporte Word: {output_path}")
            success = self.word_exporter.export_to_word(analysis_results, output_path)
            
            if success:
                print(f"   ✅ Reporte Word guardado exitosamente")
            else:
                print(f"   ❌ Error al guardar reporte Word")
            
            return success
            
        except Exception as e:
            print(f"   ❌ Error al guardar reporte Word: {e}")
            return False
    
    def save_json_report(self, analysis_results: Dict[str, Any], output_path: str):
        print(f"💾 Guardando datos JSON: {output_path}")

        json_data = self._prepare_for_json(analysis_results)

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

        print("   ✅ Datos JSON guardados exitosamente")

        
    def _generate_recommendations(
        self,
        classification,
        quality_result,
        structure_result,
        citation_analysis
    ) -> list:
        """Generate actionable recommendations based on analysis."""
        recommendations = []
        
        # Quality recommendations
        if quality_result.overall_score < 7.0:
            recommendations.append({
                'priority': 'alta',
                'message': f'La calidad general ({quality_result.overall_score:.1f}/10) necesita mejorar. '
                          'Revise las dimensiones con puntuación baja.'
            })
        
        # Check individual dimensions
        for dim_name, dim_data in quality_result.dimension_scores.items():
            if dim_data['score'] < 6.0:
                recommendations.append({
                    'priority': 'media',
                    'message': f'Dimensión "{dim_name}" tiene puntuación baja ({dim_data["score"]:.1f}). '
                              f'{dim_data.get("feedback", "Requiere atención.")}'
                })
        
        # Structure recommendations
        if not structure_result.is_valid:
            for missing in structure_result.missing_sections:
                recommendations.append({
                    'priority': 'alta',
                    'message': f'Falta la sección requerida: "{missing}". '
                              'Complete esta sección según las normas EUMIC.'
                })
        
        # Citation recommendations
        match_rate = (citation_analysis.matched_count / citation_analysis.total_citations * 100 
                     if citation_analysis.total_citations > 0 else 100)
        
        if match_rate < 90:
            recommendations.append({
                'priority': 'alta',
                'message': f'Tasa de coincidencia de citas baja ({match_rate:.1f}%). '
                          f'{citation_analysis.unmatched_count} citas no tienen referencia correspondiente.'
            })
        
        if citation_analysis.total_citations < 10:
            recommendations.append({
                'priority': 'media',
                'message': f'Número bajo de citas ({citation_analysis.total_citations}). '
                          'Considere ampliar el marco teórico con más referencias.'
            })
        
        # Classification confidence
        if classification.confidence < 0.7:
            recommendations.append({
                'priority': 'baja',
                'message': f'La clasificación tiene confianza baja ({classification.confidence:.1%}). '
                          'Verifique que el documento siga la estructura típica de su categoría.'
            })
        
        return recommendations
    
    def _format_quality_level(self, quality_level):
        return quality_level.value.replace("_", " ").title()

    def _format_category(self, article_type: ArticleType) -> str:
        """Format article type for display."""
        mapping = {
            "cientifico": "Artículo Científico",
            "academico": "Artículo Académico",
            "opinion": "Artículo de Opinión",
            "unknown": "No Clasificado"
        }

        return mapping.get(article_type.value, article_type.value.capitalize())

    def _prepare_for_json(self, data: Any) -> Any:
        if isinstance(data, dict):
            return {k: self._prepare_for_json(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [self._prepare_for_json(item) for item in data]
        elif isinstance(data, (ClassificationCategory, QualityLevel, ArticleType, ArticleSize)):
            return data.value
        else:
            return data


def main():

    """Main execution function."""
    print("\n" + "=" * 80)
    print("   SILVINA EDITORIAL ASSISTANT v0.7")
    print("   Asistente de Análisis Editorial para Documentos Académicos")
    print("=" * 80 + "\n")
    
   # Obtener ruta del documento (modo interactivo o CLI)
    if len(sys.argv) < 2:
        print("📄 SILVINA – Modo interactivo")
        document_path = input("Ingrese la ruta del documento (.docx): ").strip().strip('"')
    else:
        document_path = sys.argv[1]

      
    # Verify file exists
    if not os.path.exists(document_path):
        print(f"❌ Error: El archivo no existe: {document_path}")
        sys.exit(1)
    
    # Verify it's a .docx file
    if not document_path.lower().endswith('.docx'):
        print("❌ Error: El archivo debe ser un documento Word (.docx)")
        sys.exit(1)
    
    try:
        # Initialize Silvina
        silvina = SilvinaEditorialAssistant()
        
        # Analyze document
        results = silvina.analyze_document(document_path)
        
        # Generate output filenames
        base_name = Path(document_path).stem
        output_dir = Path(document_path).parent
        
        text_report_path = output_dir / f"{base_name}_analisis.txt"
        word_report_path = output_dir / f"{base_name}_analisis.docx"
        json_report_path = output_dir / f"{base_name}_analisis.json"
        
        # Save reports
        print("\n" + "=" * 80)
        print("📊 GENERANDO REPORTES")
        print("=" * 80 + "\n")
        
        silvina.save_text_report(results, str(text_report_path))
        silvina.save_word_report(results, str(word_report_path))
        silvina.save_json_report(results, str(json_report_path))
        
        # Print summary
        # Print summary
        print("\n" + "=" * 80)
        print("🎉 ANÁLISIS COMPLETADO")
        print("=" * 80)
        print("\nRESUMEN:")
        print(f"  📄 Documento analizado: {results['filename']}")
        print(f"  📝 Total de palabras: {results['document_info']['word_count']:,}")
        
        print(f"  📝 Total de palabras: {results['document_info']['word_count']:,}")
        print(f"  📋 Total de párrafos: {results['document_info']['paragraph_count']}")
        print(f"\n  🏷️  Tipo: {results['classification']['category'].value.upper()}")
               
        print(f"  📏 Tamaño: {results['classification']['article_size'].value.upper()}")
        print(f"  💭 Razonamiento: {results['classification']['reasoning']}")

        print(f"\n  ⭐ ANÁLISIS DE CALIDAD: {results['quality_analysis']['overall_score']:.1f}/10")
        for dim, data in results['quality_analysis']['dimensions'].items():
            print(f"     • {dim.capitalize()}: {data['score']:.1f}/10 - {data['feedback']}")

        print(f"\n  📋 ESTRUCTURA: {'✓ Válida' if results['structure_validation']['is_valid'] else '✗ Incompleta'}")
        if results['structure_validation']['missing_sections']:
            print("     Missing sections:")
            for sec in results['structure_validation']['missing_sections']:
                print(f"       - {sec}")

        print(f"\n  📚 CITAS: {results['citations_analysis']['total_citations']} detectadas")

        print(f"\n  💡 ANÁLISIS FINAL:")
        for rec in results['recommendations']:
            color = {'critico': '🔴', 'moderado': '🟢', 'aceptable': '🟡', 'opcional': '🔵'}.get(rec['priority'], '⚪')
            print(f"     {color} {rec['priority'].upper()}: {rec['message']}")

        print(f"\n  💾 Reportes: {output_dir}")
        print("=" * 80 + "\n")
                
                
    except KeyboardInterrupt:
        print("\n\n⚠ Análisis interrumpido por el usuario")
        sys.exit(0)
    except Exception as e:
        print(f"\n\n❌ Error fatal: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()