"""
word_exporter.py
Export analysis results to professionally formatted Word (.docx) document.
PROFESSIONAL VERSION - Tables, colors, proper formatting
"""
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from typing import Dict, Any
from datetime import datetime


class WordExporter:
    """Export analysis results to professionally formatted Word document."""

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Ejecute: pip install python-docx")

    def export_to_word(self, analysis_results: Dict[str, Any], output_path: str) -> bool:
        try:
            doc = Document()

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(12)

            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = 1.15
            paragraph_format.space_after = Pt(0)
            paragraph_format.space_before = Pt(0)

            import os
            script_dir = os.path.dirname(os.path.abspath(__file__))
            logo_path = os.path.join(script_dir, '..', 'assets', 'logo.jpg')

            if os.path.exists(logo_path):
                self._add_header_logo(doc, logo_path)
            else:
                print(f"      ⚠️  Logo not found at: {logo_path}")

            self._add_page_numbers(doc)
            self._add_title_page(doc, analysis_results)
            self._add_executive_summary(doc, analysis_results)
            self._add_document_info(doc, analysis_results)
            self._add_classification(doc, analysis_results)
            self._add_quality_analysis(doc, analysis_results)
            self._add_grammar_analysis(doc, analysis_results)
            self._add_apa_validation(doc, analysis_results)
            self._add_structure_validation(doc, analysis_results)
            self._add_citations_analysis(doc, analysis_results)
            self._add_recommendations(doc, analysis_results)
            self._add_footer(doc)

            doc.save(output_path)
            return True

        except Exception as e:
            print(f"      ❌ Error al crear documento Word: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _add_title_page(self, doc, results):
        title = doc.add_heading('INFORME DE ANÁLISIS EDITORIAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.bold = True
            run.font.underline = True

        doc.add_paragraph()

        doc_name = doc.add_paragraph()
        doc_name.add_run(results['document_info'].get('title', results['filename'])).bold = True
        doc_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_name.runs[0].font.size = Pt(12)

    def _add_header_logo(self, doc, logo_path: str):
        try:
            import os
            if not os.path.exists(logo_path):
                print(f"      ⚠️ Logo no encontrado")
                return

            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].clear()

            table = header.add_table(rows=1, cols=2, width=Inches(6.5))

            left_cell = table.rows[0].cells[0]
            left_cell.width = Inches(4.5)
            left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = left_para.add_run('Generado por Silvina Revisor Editorial 0.9\n')
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

            date_run = left_para.add_run(datetime.now().strftime('%d de %B de %Y'))
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(128, 128, 128)

            right_cell = table.rows[0].cells[1]
            right_cell.width = Inches(2.0)
            right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            logo_run = right_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.8))

            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            print(f"      ✅ Logo agregado exitosamente")

        except Exception as e:
            print(f"      ⚠️ Error al agregar logo: {e}")

    def _add_page_numbers(self, doc):
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            section = doc.sections[0]
            footer = section.footer
            footer.paragraphs[0].clear()

            para = footer.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run()

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            run._element.append(instrText)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._element.append(fldChar2)

            run = para.add_run(' de ')
            run.font.size = Pt(10)

            run = para.add_run()

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar3)

            instrText2 = OxmlElement('w:instrText')
            instrText2.set(qn('xml:space'), 'preserve')
            instrText2.text = 'NUMPAGES'
            run._element.append(instrText2)

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            run._element.append(fldChar4)

            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

            print(f"      ✅ Números de página agregados")

        except Exception as e:
            print(f"      ⚠️ Error al agregar números de página: {e}")

    def _add_executive_summary(self, doc, results):
        heading = doc.add_heading('RESUMEN EJECUTIVO', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        can_publish, reason = self._determine_publishability(results)

        decision_para = doc.add_paragraph()
        decision_run = decision_para.add_run(
            "✅ APTO PARA PUBLICACIÓN" if can_publish
            else "⚠️ REQUIERE REVISIÓN"
        )
        decision_run.bold = True
        decision_run.font.size = Pt(16)
        decision_run.font.color.rgb = RGBColor(0, 128, 0) if can_publish else RGBColor(192, 0, 0)

        doc.add_paragraph(reason)
        doc.add_paragraph()

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        quality = results['quality_analysis']
        citations = results['citations_analysis']

        table.rows[1].cells[0].text = 'Calidad General'
        table.rows[1].cells[1].text = f"{quality['overall_score']:.1f}/10"

        table.rows[2].cells[0].text = 'Gramática y Ortografía'
        table.rows[2].cells[1].text = f"{quality['gramatica']['score']:.1f}/10"

        table.rows[3].cells[0].text = 'Estructura'
        table.rows[3].cells[1].text = "✓ Válida" if results['structure_validation']['is_valid'] else "✗ Incompleta"

        table.rows[4].cells[0].text = 'Errores APA 7'
        apa_count = citations.get('apa_violations', 0)
        table.rows[4].cells[1].text = f"{apa_count} detectados" if apa_count > 0 else "Sin errores"

        table.rows[5].cells[0].text = 'Tasa de Coincidencia'
        if citations['total_citations'] > 0:
            match_rate = citations['matched_count'] / citations['total_citations'] * 100
            table.rows[5].cells[1].text = f"{match_rate:.1f}%"
        else:
            table.rows[5].cells[1].text = "N/A"

        doc.add_paragraph()

    def _add_document_info(self, doc, results):
        heading = doc.add_heading('📄 INFORMACIÓN DEL DOCUMENTO', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        info = results['document_info']

        p = doc.add_paragraph()
        p.add_run('Título: ').bold = True
        p.add_run(info.get('title', 'No especificado'))

        p = doc.add_paragraph()
        p.add_run('Autor(es): ').bold = True
        p.add_run(info.get('authors', 'No identificado'))

        p = doc.add_paragraph()
        p.add_run('Total de palabras: ').bold = True
        p.add_run(f"{info['word_count']:,}")

        p = doc.add_paragraph()
        p.add_run('Total de caracteres: ').bold = True
        p.add_run(f"{info['char_count']:,}")

        p = doc.add_paragraph()
        p.add_run('Páginas estimadas: ').bold = True
        p.add_run(str(info['estimated_pages']))

    def _add_classification(self, doc, results):
        heading = doc.add_heading('🏷️ CLASIFICACIÓN DEL ARTÍCULO', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        classification = results['classification']

        p = doc.add_paragraph()
        p.add_run('Categoría: ').bold = True
        p.add_run(classification['category'].value.upper())

        p = doc.add_paragraph()
        p.add_run('Confianza: ').bold = True
        conf = f"{classification['confidence']:.1%}" if classification['confidence'] is not None else "—"
        p.add_run(conf)
        
        if classification.get('reasoning'):
            p = doc.add_paragraph()
            p.add_run('Razonamiento: ').bold = True
            p.add_run(classification['reasoning'])

    def _add_quality_analysis(self, doc, results):
        heading = doc.add_heading('⭐ ANÁLISIS DE CALIDAD SEMÁNTICA', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        quality = results['quality_analysis']

        p = doc.add_paragraph()
        p.add_run('Puntuación General: ').bold = True
        score_run = p.add_run(f"{quality['overall_score']:.1f}/10")
        score_run.bold = True
        score_run.font.size = Pt(14)
        if quality['overall_score'] >= 8:
            score_run.font.color.rgb = RGBColor(0, 128, 0)
        elif quality['overall_score'] >= 6:
            score_run.font.color.rgb = RGBColor(255, 140, 0)
        else:
            score_run.font.color.rgb = RGBColor(192, 0, 0)

        if quality.get('dimensions'):
            for dim_name, dim_data in quality['dimensions'].items():
                doc.add_heading(dim_name.capitalize(), level=3)

                p = doc.add_paragraph()
                p.add_run('Puntuación: ').bold = True
                p.add_run(f"{dim_data['score']:.1f}/10")

                if dim_data.get('feedback'):
                    doc.add_paragraph(dim_data['feedback'])

    def _add_grammar_analysis(self, doc, results):
        heading = doc.add_heading('📝 GRAMÁTICA Y ORTOGRAFÍA', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        grammar = results['quality_analysis']['gramatica']

        p = doc.add_paragraph()
        p.add_run('Puntuación: ').bold = True
        p.add_run(f"{grammar['score']:.1f}/10")

        p = doc.add_paragraph()
        p.add_run('Estado: ').bold = True
        p.add_run(grammar['feedback'])

        if grammar.get('errors') and len(grammar['errors']) > 0:
            doc.add_paragraph()
            doc.add_paragraph('Errores Detectados:').bold = True

            for err in grammar['errors'][:5]:
                doc.add_paragraph(err['message'], style='List Number')

                context_text = err['context'] if len(err['context']) < 150 else err['context'][:150] + "..."
                doc.add_paragraph(f"   Contexto: \"{context_text}\"", style='List Bullet 2')

                if err.get('replacements'):
                    doc.add_paragraph(f"   Sugerencia: {', '.join(err['replacements'][:3])}", style='List Bullet 2')

    def _add_apa_validation(self, doc, results):
        heading = doc.add_heading('📖 VALIDACIÓN APA 7 (ESPAÑOL)', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        apa = results.get('apa_validation', {})
        violations = apa.get('violations', [])

        if len(violations) == 0:
            p = doc.add_paragraph()
            total_citations = results['citations_analysis'].get('total_citations', 0)
            if total_citations == 0:
                p.add_run('ℹ️ Sin citaciones en texto detectadas — no se pudo validar formato APA 7').font.color.rgb = RGBColor(128, 128, 128)
            else:
                p.add_run('✅ Sin errores de formato APA 7 detectados').font.color.rgb = RGBColor(0, 128, 0)

            doc.add_paragraph()
        else:
            from collections import defaultdict
            by_type = defaultdict(list)
            for v in violations:
                by_type[v['error_type']].append(v)

            for error_type, errors in by_type.items():
                doc.add_paragraph(f"CITACIÓN INCORRECTA ({len(errors)}):", style='Heading 3')

                for i, err in enumerate(errors[:5], 1):
                    doc.add_paragraph(f"{i}. Citación: {err['citation']}")
                    doc.add_paragraph(f"   Ubicación: \"{err['location']}\"", style='List Bullet 2')
                    doc.add_paragraph(f"   Problema: {err['explanation']}", style='List Bullet 2')
                    if err.get('correction'):
                        doc.add_paragraph(f"   Corrección: {err['correction']}", style='List Bullet 2')

    def _add_structure_validation(self, doc, results):
        heading = doc.add_heading('📋 VALIDACIÓN DE ESTRUCTURA (EUMIC)', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        structure = results['structure_validation']

        if structure['is_valid']:
            p = doc.add_paragraph()
            p.add_run('✓ Estructura válida según normas EUMIC').font.color.rgb = RGBColor(0, 128, 0)
        else:
            p = doc.add_paragraph()
            p.add_run('✗ Estructura incompleta').font.color.rgb = RGBColor(192, 0, 0)

            if structure.get('missing_sections'):
                doc.add_paragraph()
                doc.add_paragraph('Secciones faltantes:').bold = True
                for section in structure['missing_sections']:
                    doc.add_paragraph(f"• {section}", style='List Bullet')

    def _add_citations_analysis(self, doc, results):
        heading = doc.add_heading('📚 ANÁLISIS DE CITAS Y REFERENCIAS', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        citations = results['citations_analysis']

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Total de citas en texto'
        table.rows[0].cells[1].text = str(citations['total_citations'])

        table.rows[1].cells[0].text = 'Total de referencias bibliográficas'
        table.rows[1].cells[1].text = str(citations['total_references'])

        table.rows[2].cells[0].text = 'Citas con referencia'
        table.rows[2].cells[1].text = str(citations['matched_count'])

        table.rows[3].cells[0].text = 'Tasa de coincidencia'
        if citations['total_citations'] > 0:
            rate = citations['matched_count'] / citations['total_citations'] * 100
            table.rows[3].cells[1].text = f"{rate:.1f}%"
        else:
            table.rows[3].cells[1].text = "N/A"

    def _add_recommendations(self, doc, results):
        if not results.get('recommendations'):
            return

        heading = doc.add_heading('💡 RECOMENDACIONES', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        recommendations = results['recommendations']

        final_rec = [r for r in recommendations if r.get('priority') in ['critica', 'advertencia', 'aprobado']]
        if final_rec:
            rec = final_rec[0]
            p = doc.add_paragraph()
            p.add_run(rec['message']).bold = True
            p.runs[0].font.size = Pt(12)
            if rec['priority'] == 'critica':
                p.runs[0].font.color.rgb = RGBColor(192, 0, 0)
            elif rec['priority'] == 'advertencia':
                p.runs[0].font.color.rgb = RGBColor(192, 0, 0)
            else:
                p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

        others = [r for r in recommendations if r.get('priority') not in ['critica', 'advertencia', 'aprobado']]
        if others:
            doc.add_paragraph('Recomendaciones específicas:').bold = True
            for rec in others:
                priority_icon = {'alta': '🔴', 'media': '🟡', 'baja': '🟢'}.get(rec['priority'], '⚪')
                doc.add_paragraph(f"{priority_icon} {rec['message']}", style='List Bullet')

    def _add_footer(self, doc):
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)

        footer = doc.add_paragraph()
        footer.add_run("Generado por Silvina Editorial Assistant v0.9 | ").italic = True
        footer.add_run(datetime.now().strftime('%d/%m/%Y %H:%M')).italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        footer.runs[1].font.size = Pt(9)
        footer.runs[1].font.color.rgb = RGBColor(128, 128, 128)

    def _determine_publishability(self, results: Dict[str, Any]) -> tuple:
        """Determine if document can be published."""
        quality_score = results['quality_analysis']['overall_score']
        grammar_score = results['quality_analysis']['gramatica']['score']
        is_structure_valid = results['structure_validation']['is_valid']
        apa_violations = results['citations_analysis'].get('apa_violations', 0)
        total_citations = results['citations_analysis'].get('total_citations', 0)

        if quality_score >= 7.0 and grammar_score >= 7.0 and is_structure_valid and apa_violations == 0 and total_citations > 0:
            return (True, "El documento cumple con todos los estándares de calidad, estructura y formato APA 7 requeridos por las normas EUMIC.")
        elif total_citations == 0:
            return (False, "No se detectaron citas APA en el texto. Verifique el formato de citación según normas APA 7.")
        elif apa_violations > 0 or grammar_score < 7.0 or quality_score < 7.0:
            return (False, f"El documento requiere revisión. Calidad: {quality_score:.1f}/10, Gramática: {grammar_score:.1f}/10, Errores APA: {apa_violations}.")
        elif not is_structure_valid:
            return (False, "Estructura incompleta según normas EUMIC. Complete las secciones faltantes.")
        else:
            return (False, "El documento requiere mejoras antes de la publicación.")
