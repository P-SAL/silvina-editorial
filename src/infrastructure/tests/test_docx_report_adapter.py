import tempfile
from pathlib import Path
from unittest import TestCase

from docx import Document

from src.domain.dtos.editorial_suitability_dto import EditorialSuitabilityDTO
from src.domain.dtos.quality_result_dto import QualityResultDTO
from src.domain.enums.quality_level import QualityLevel
from src.infrastructure.adapters.report.docx_report_adapter import DocxReportAdapter
from src.infrastructure.tests.adapters.report.fixtures import ReportFixtures


def _all_paragraph_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


class TestDocxReportAdapterEditorialSuitability(TestCase):
    """Integration test: real Document generation and Word-file round-trip."""

    def setUp(self):
        self._tmp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self._tmp_dir.cleanup)
        self.output_path = str(Path(self._tmp_dir.name) / "report.docx")
        self.adapter = DocxReportAdapter(logo_path=None, settings=ReportFixtures.make_settings())

    def test_export_renders_editorial_suitability_section_when_present(self):
        suitability = EditorialSuitabilityDTO(
            contribution_verdict="SUSTENTADA",
            contribution_phrase="Propone un marco de análisis original.",
            contribution_observation=(
                "Contribución sustentada — Propone un marco de análisis original."
            ),
            alignment_verdict="ALINEADO",
            alignment_lines="Línea 1 y 2.",
            alignment_justification="Se relaciona directamente con las líneas mencionadas.",
        )
        quality = QualityResultDTO(
            overall_score=8.0,
            quality_level=QualityLevel.GOOD,
            editorial_suitability=suitability,
        )
        report_input = ReportFixtures.make_report_input_dto(quality=quality)

        self.adapter.export(report_input=report_input, output_path=self.output_path)

        document = Document(self.output_path)
        full_text = _all_paragraph_text(document)
        self.assertIn("SUSTENTADA", full_text)
        self.assertIn("Propone un marco de análisis original.", full_text)
        self.assertIn("ALINEADO", full_text)
        self.assertIn("Línea 1 y 2.", full_text)
        self.assertIn("Se relaciona directamente con las líneas mencionadas.", full_text)

    def test_export_omits_editorial_suitability_section_when_absent(self):
        quality = QualityResultDTO(overall_score=8.0, quality_level=QualityLevel.GOOD)
        report_input = ReportFixtures.make_report_input_dto(quality=quality)

        self.adapter.export(report_input=report_input, output_path=self.output_path)

        document = Document(self.output_path)
        full_text = _all_paragraph_text(document)
        self.assertNotIn("SUSTENTADA", full_text)
        self.assertNotIn("ALINEADO", full_text)
