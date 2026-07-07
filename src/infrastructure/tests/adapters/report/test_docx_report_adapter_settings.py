from unittest import TestCase

from docx import Document

from src.domain.dtos.apa_violation_dto import ApaViolationDTO
from src.domain.dtos.grammar_error_dto import GrammarErrorDTO
from src.domain.enums.apa_error_type import ApaErrorType
from src.infrastructure.adapters.report.docx_report_adapter import DocxReportAdapter
from src.infrastructure.tests.adapters.report.fixtures import ReportFixtures


class TestDocxReportAdapterSettings(TestCase):
    def test_estimated_pages_uses_settings_words_per_page(self):
        settings = ReportFixtures.make_settings(words_per_page=100)
        adapter = DocxReportAdapter(logo_path=None, settings=settings)
        doc = Document()
        doc_content = ReportFixtures.make_doc_content_mock(word_count=250)
        report_input = ReportFixtures.make_report_input_dto(document_content=doc_content)

        adapter._add_document_info(doc=doc, report_input=report_input)

        pages_text = [p.text for p in doc.paragraphs if "Páginas estimadas" in p.text]
        self.assertEqual(pages_text[0], "Páginas estimadas: 2")

    def test_grammar_errors_limited_by_max_errors_displayed(self):
        settings = ReportFixtures.make_settings(
            max_errors_displayed=1, context_truncation_limit=150, max_replacements=3
        )
        adapter = DocxReportAdapter(logo_path=None, settings=settings)
        doc = Document()
        errors = [
            GrammarErrorDTO(
                number=i, message=f"error {i}", context="ctx", offset=0, length=1, replacements=[]
            )
            for i in range(3)
        ]
        grammar = ReportFixtures.make_grammar_mock(errors=errors)
        report_input = ReportFixtures.make_report_input_dto(grammar=grammar)

        adapter._add_grammar_analysis(doc=doc, report_input=report_input)

        error_paragraphs = [p.text for p in doc.paragraphs if p.text.startswith("error")]
        self.assertEqual(len(error_paragraphs), 1)

    def test_grammar_context_truncated_by_context_truncation_limit(self):
        settings = ReportFixtures.make_settings(
            max_errors_displayed=5, context_truncation_limit=10, max_replacements=3
        )
        adapter = DocxReportAdapter(logo_path=None, settings=settings)
        doc = Document()
        errors = [
            GrammarErrorDTO(
                number=1,
                message="error",
                context="x" * 50,
                offset=0,
                length=1,
                replacements=[],
            )
        ]
        grammar = ReportFixtures.make_grammar_mock(errors=errors)
        report_input = ReportFixtures.make_report_input_dto(grammar=grammar)

        adapter._add_grammar_analysis(doc=doc, report_input=report_input)

        context_paragraphs = [p.text for p in doc.paragraphs if "Contexto" in p.text]
        self.assertEqual(context_paragraphs[0], '   Contexto: "' + "x" * 10 + '..."')

    def test_grammar_replacements_limited_by_max_replacements(self):
        settings = ReportFixtures.make_settings(
            max_errors_displayed=5, context_truncation_limit=150, max_replacements=1
        )
        adapter = DocxReportAdapter(logo_path=None, settings=settings)
        doc = Document()
        errors = [
            GrammarErrorDTO(
                number=1,
                message="error",
                context="ctx",
                offset=0,
                length=1,
                replacements=["a", "b", "c"],
            )
        ]
        grammar = ReportFixtures.make_grammar_mock(errors=errors)
        report_input = ReportFixtures.make_report_input_dto(grammar=grammar)

        adapter._add_grammar_analysis(doc=doc, report_input=report_input)

        suggestion_paragraphs = [p.text for p in doc.paragraphs if "Sugerencia" in p.text]
        self.assertEqual(suggestion_paragraphs[0], "   Sugerencia: a")

    def test_apa_violations_limited_by_max_errors_displayed(self):
        settings = ReportFixtures.make_settings(max_errors_displayed=1)
        adapter = DocxReportAdapter(logo_path=None, settings=settings)
        doc = Document()
        violations = [
            ApaViolationDTO(
                citation_text=f"cite{i}",
                error_type=ApaErrorType.CONJUNCTION_ERROR,
                location=i,
                explanation="x",
                correction="y",
            )
            for i in range(3)
        ]
        apa = ReportFixtures.make_apa_validation_mock(violations=violations)
        report_input = ReportFixtures.make_report_input_dto(apa_validation=apa)

        adapter._add_apa_validation(doc=doc, report_input=report_input)

        citation_paragraphs = [p.text for p in doc.paragraphs if "Citación:" in p.text]
        self.assertEqual(len(citation_paragraphs), 1)
