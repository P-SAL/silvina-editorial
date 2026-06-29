from unittest import TestCase
from unittest.mock import MagicMock, patch

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from src.domain.enums.severity_level import SeverityLevel
from src.infrastructure.adapters.document.docx_eumic_adapter import DocxEumicAdapter


def _mock_length(cm_value: float) -> MagicMock:
    """Return a mock Length object using the project twips formula."""
    mock = MagicMock()
    mock.cm = cm_value
    mock.twips = Cm(cm_value).twips
    return mock


def _make_mock_section(margin_cm: float = 2.5) -> MagicMock:
    section = MagicMock()
    section.top_margin = _mock_length(margin_cm)
    section.bottom_margin = _mock_length(margin_cm)
    section.left_margin = _mock_length(margin_cm)
    section.right_margin = _mock_length(margin_cm)
    return section


def _make_run(
    xml: str = "<w:r/>",
    font_name: str | None = None,
    font_size_pt: float | None = None,
) -> MagicMock:
    run = MagicMock()
    run._element.xml = xml
    run.font.name = font_name
    if font_size_pt is not None:
        size_mock = MagicMock()
        size_mock.pt = font_size_pt
        run.font.size = size_mock
    else:
        run.font.size = None
    return run


def _make_paragraph(
    text: str = "",
    alignment: object = WD_ALIGN_PARAGRAPH.JUSTIFY,
    runs: list | None = None,
) -> MagicMock:
    para = MagicMock()
    para.text = text
    para.alignment = alignment
    para.runs = runs if runs is not None else []
    return para


def _make_formula_run() -> MagicMock:
    """Return a run whose XML contains an oMath element."""
    run = MagicMock()
    run._element.xml = "<m:oMath>x + y = z</m:oMath>"
    run.font.name = "Times New Roman"
    run.font.size = None
    return run


def _make_image_rel() -> MagicMock:
    rel = MagicMock()
    rel.target_ref = "word/media/image1.png"
    return rel


def _make_compliant_doc() -> MagicMock:
    """Build a minimal Document mock that produces zero violations for word_count < 1000."""
    doc = MagicMock()
    doc.sections = [_make_mock_section(2.5)]
    doc.paragraphs = []
    doc.tables = []
    doc.part.rels = {}
    return doc


class TestDocxEumicAdapterFullyCompliantDocument(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_empty_list_when_fully_compliant_document(self, mock_document_class):
        mock_document_class.return_value = _make_compliant_doc()
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        self.assertEqual(result, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_list_instance(self, mock_document_class):
        mock_document_class.return_value = _make_compliant_doc()
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        self.assertIsInstance(result, list)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_calls_document_with_provided_path(self, mock_document_class):
        mock_document_class.return_value = _make_compliant_doc()
        DocxEumicAdapter().inspect(docx_path="/path/to/doc.docx", word_count=500)
        mock_document_class.assert_called_once_with("/path/to/doc.docx")


class TestDocxEumicAdapterMarginCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_margin_violation_when_top_margin_is_too_small(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        section = _make_mock_section(2.5)
        section.top_margin = _mock_length(1.0)
        doc.sections = [section]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        messages = [violation.message for violation in result]
        self.assertTrue(any("Margen superior" in message for message in messages))

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_no_margin_violations_when_sections_list_is_empty(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.sections = []
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        margin_violations = [violation for violation in result if "Margen" in violation.message]
        self.assertEqual(margin_violations, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_four_margin_violations_when_all_margins_are_non_compliant(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.sections = [_make_mock_section(1.0)]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        margin_violations = [violation for violation in result if "Margen" in violation.message]
        self.assertEqual(len(margin_violations), 4)
        self.assertTrue(all(v.severity == SeverityLevel.WARNING for v in margin_violations))


class TestDocxEumicAdapterFontCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_font_violation_when_non_standard_font_is_used(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        run = _make_run(font_name="Comic Sans MS")
        doc.paragraphs = [_make_paragraph(text="Content", runs=[run])]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        font_violations = [
            violation for violation in result if "Fuentes no estándar" in violation.message
        ]
        self.assertEqual(len(font_violations), 1)
        self.assertIn("Comic Sans MS", font_violations[0].details)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_no_font_violation_when_all_fonts_are_standard(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        run = _make_run(font_name="Times New Roman")
        doc.paragraphs = [_make_paragraph(text="Content", runs=[run])]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        font_violations = [
            violation for violation in result if "Fuentes no estándar" in violation.message
        ]
        self.assertEqual(font_violations, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_font_size_violation_when_non_standard_size_is_detected(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        run = _make_run(font_name="Times New Roman", font_size_pt=8.0)
        doc.paragraphs = [_make_paragraph(text="Content", runs=[run])]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        size_violations = [
            violation for violation in result if "Tamaños de fuente" in violation.message
        ]
        self.assertEqual(len(size_violations), 1)
        self.assertEqual(size_violations[0].severity, SeverityLevel.INFO)


class TestDocxEumicAdapterAlignmentCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_alignment_violation_when_majority_of_paragraphs_are_not_justified(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.paragraphs = [
            _make_paragraph(text="First", alignment=WD_ALIGN_PARAGRAPH.LEFT),
            _make_paragraph(text="Second", alignment=WD_ALIGN_PARAGRAPH.LEFT),
            _make_paragraph(text="Third", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY),
            _make_paragraph(text="Fourth", alignment=WD_ALIGN_PARAGRAPH.LEFT),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        alignment_violations = [
            violation for violation in result if "justificado" in violation.message
        ]
        self.assertEqual(len(alignment_violations), 1)
        self.assertEqual(alignment_violations[0].severity, SeverityLevel.WARNING)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_no_alignment_violation_when_all_paragraphs_are_justified(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.paragraphs = [
            _make_paragraph(text="First", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY),
            _make_paragraph(text="Second", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        alignment_violations = [
            violation for violation in result if "justificado" in violation.message
        ]
        self.assertEqual(alignment_violations, [])


class TestDocxEumicAdapterFigureCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_no_figure_violations_when_no_images_in_document(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.part.rels = {}
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        figure_violations = [violation for violation in result if violation.category == "Figuras"]
        self.assertEqual(figure_violations, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_caption_violation_when_image_has_no_caption(self, mock_document_class):
        doc = _make_compliant_doc()
        doc.part.rels = {"rId1": _make_image_rel()}
        doc.paragraphs = []
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        caption_violations = [
            violation for violation in result if "título formal" in violation.message
        ]
        self.assertEqual(len(caption_violations), 1)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_numbering_violation_when_figure_captions_are_out_of_order(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.part.rels = {"rId1": _make_image_rel(), "rId2": _make_image_rel()}
        doc.paragraphs = [
            _make_paragraph(text="Figura 2. First image"),
            _make_paragraph(text="Figura 1. Second image"),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        numbering_violations = [
            violation for violation in result if "Numeración de figuras" in violation.message
        ]
        self.assertEqual(len(numbering_violations), 1)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_suppresses_attribute_error_from_doc_part_rels_access(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        rels_mock = MagicMock()
        rels_mock.values.side_effect = AttributeError("no rels")
        doc.part.rels = rels_mock
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        figure_violations = [violation for violation in result if violation.category == "Figuras"]
        self.assertEqual(figure_violations, [])


class TestDocxEumicAdapterTableCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_no_table_violations_when_no_tables_in_document(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.tables = []
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        table_violations = [violation for violation in result if violation.category == "Tablas"]
        self.assertEqual(table_violations, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_title_violation_when_tables_have_no_titles(self, mock_document_class):
        doc = _make_compliant_doc()
        doc.tables = [MagicMock(), MagicMock()]
        doc.paragraphs = []
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        title_violations = [
            violation for violation in result if "sin título descriptivo" in violation.message
        ]
        self.assertEqual(len(title_violations), 1)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_numbering_violation_when_table_titles_are_out_of_order(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.tables = [MagicMock(), MagicMock()]
        doc.paragraphs = [
            _make_paragraph(text="Tabla 2. First table"),
            _make_paragraph(text="Tabla 1. Second table"),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        numbering_violations = [
            violation for violation in result if "Numeración de tablas" in violation.message
        ]
        self.assertEqual(len(numbering_violations), 1)


class TestDocxEumicAdapterFormulaCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_no_formula_violations_when_no_formulas_in_document(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        run = _make_run(xml="<w:r/>")
        doc.paragraphs = [_make_paragraph(text="No formulas here", runs=[run])]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        formula_violations = [violation for violation in result if violation.category == "Fórmulas"]
        self.assertEqual(formula_violations, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_alignment_violation_when_formula_paragraph_is_not_centered(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        formula_para = _make_paragraph(
            text="",
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            runs=[_make_formula_run()],
        )
        doc.paragraphs = [formula_para]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        formula_violations = [
            violation for violation in result if "Fórmulas no centradas" in violation.message
        ]
        self.assertEqual(len(formula_violations), 1)
        self.assertEqual(formula_violations[0].severity, SeverityLevel.INFO)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_suppresses_attribute_error_when_run_element_xml_raises(
        self, mock_document_class
    ):
        class _ElementWithNoXml:
            @property
            def xml(self) -> str:
                raise AttributeError("no xml attribute")

        bad_run = MagicMock()
        bad_run._element = _ElementWithNoXml()
        bad_run.font.name = None
        bad_run.font.size = None
        doc = _make_compliant_doc()
        doc.paragraphs = [_make_paragraph(text="formula paragraph", runs=[bad_run])]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=500)
        self.assertIsInstance(result, list)


class TestDocxEumicAdapterAbstractKeywordCheck(TestCase):
    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_skips_abstract_check_when_word_count_is_below_minimum(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=999)
        abstract_violations = [
            violation for violation in result if violation.category == "Resumen y Palabras Clave"
        ]
        self.assertEqual(abstract_violations, [])

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_missing_abstract_violation_when_no_abstract_section_found(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.paragraphs = [_make_paragraph(text="Palabras clave: investigación, ciencia, datos")]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=2000)
        abstract_violations = [
            violation for violation in result if "Resumen/Abstract" in violation.message
        ]
        self.assertEqual(len(abstract_violations), 1)
        self.assertEqual(abstract_violations[0].severity, SeverityLevel.CRITICAL)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_abstract_length_violation_when_abstract_content_is_too_short(
        self, mock_document_class
    ):
        doc = _make_compliant_doc()
        doc.paragraphs = [
            _make_paragraph(text="Resumen"),
            _make_paragraph(text="Short abstract."),
            _make_paragraph(text="Palabras clave: investigación, ciencia, datos"),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=2000)
        length_violations = [
            violation for violation in result if "Extensión del resumen" in violation.message
        ]
        self.assertEqual(len(length_violations), 1)
        self.assertEqual(length_violations[0].severity, SeverityLevel.WARNING)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_missing_keywords_violation_when_no_keywords_section_found(
        self, mock_document_class
    ):
        abstract_content = " ".join(["word"] * 150)
        doc = _make_compliant_doc()
        doc.paragraphs = [
            _make_paragraph(text="Resumen"),
            _make_paragraph(text=abstract_content),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=2000)
        keyword_violations = [
            violation for violation in result if "palabras clave" in violation.message.lower()
        ]
        self.assertEqual(len(keyword_violations), 1)
        self.assertEqual(keyword_violations[0].severity, SeverityLevel.CRITICAL)

    @patch("src.infrastructure.adapters.document.docx_eumic_adapter.Document")
    def test_inspect_returns_keyword_count_violation_when_keyword_count_exceeds_maximum(
        self, mock_document_class
    ):
        abstract_content = " ".join(["word"] * 150)
        doc = _make_compliant_doc()
        doc.paragraphs = [
            _make_paragraph(text="Resumen"),
            _make_paragraph(text=abstract_content),
            _make_paragraph(text="Palabras clave: alpha, beta, gamma, delta, epsilon, zeta"),
        ]
        mock_document_class.return_value = doc
        result = DocxEumicAdapter().inspect(docx_path="test.docx", word_count=2000)
        keyword_count_violations = [
            violation for violation in result if "Número incorrecto" in violation.message
        ]
        self.assertEqual(len(keyword_count_violations), 1)
        self.assertEqual(keyword_count_violations[0].severity, SeverityLevel.WARNING)
