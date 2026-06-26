from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase

from docx import Document

from src.domain.exceptions.document_errors import DocumentNotFound, DocumentUnreadable
from src.infrastructure.adapters.document.docx_text_adapter import DocxTextAdapter

SAMPLE_DOCUMENT = (
    Path(__file__).parent.parent.parent.parent.parent.parent
    / "docs"
    / "sample-documents"
    / "1. test_Científico.docx"
)


class TestDocxTextAdapter(TestCase):
    def setUp(self):
        self.adapter = DocxTextAdapter()

    def test_strips_and_filters_empty_paragraphs(self):
        paragraphs = self.adapter.read_paragraphs(str(SAMPLE_DOCUMENT))
        self.assertTrue(all(paragraph == paragraph.strip() for paragraph in paragraphs))
        self.assertTrue(all(paragraph != "" for paragraph in paragraphs))

    def test_preserves_paragraph_order(self):
        with TemporaryDirectory() as temporary_directory:
            document_path = Path(temporary_directory) / "ordered.docx"
            document = Document()
            document.add_paragraph("First")
            document.add_paragraph("Second")
            document.add_paragraph("Third")
            document.save(str(document_path))

            paragraphs = self.adapter.read_paragraphs(str(document_path))

            self.assertEqual(paragraphs, ["First", "Second", "Third"])

    def test_no_non_empty_paragraphs_returns_empty_list(self):
        with TemporaryDirectory() as temporary_directory:
            document_path = Path(temporary_directory) / "blank.docx"
            document = Document()
            document.add_paragraph("   ")
            document.add_paragraph("")
            document.save(str(document_path))

            paragraphs = self.adapter.read_paragraphs(str(document_path))

            self.assertEqual(paragraphs, [])

    def test_missing_file_raises_document_not_found(self):
        with self.assertRaises(DocumentNotFound):
            self.adapter.read_paragraphs("nonexistent-file.docx")

    def test_corrupt_file_raises_document_unreadable(self):
        with TemporaryDirectory() as temporary_directory:
            document_path = Path(temporary_directory) / "corrupt.docx"
            document_path.write_bytes(b"not a valid docx package")

            with self.assertRaises(DocumentUnreadable):
                self.adapter.read_paragraphs(str(document_path))

    def test_valid_file_raises_no_exception(self):
        self.adapter.read_paragraphs(str(SAMPLE_DOCUMENT))
