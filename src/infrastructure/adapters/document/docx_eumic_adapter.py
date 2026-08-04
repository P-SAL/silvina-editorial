from re import search, split

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from src.domain.document.document_format_inspection_port import DocumentFormatInspectionPort
from src.domain.dtos.eumic_violation_dto import EumicViolationDTO
from src.domain.enums.allowed_font import AllowedFont
from src.domain.enums.formula_xml_marker import FormulaXmlMarker
from src.infrastructure.adapters.document.eumic_document_standards import (
    ABSTRACT_MAX_WORD_COUNT,
    ABSTRACT_MIN_WORD_COUNT,
    ABSTRACT_PARAGRAPH_LOOKAHEAD,
    ABSTRACT_SECTION_KEYWORDS,
    FIGURE_CAPTION_PREFIXES,
    FIGURE_NUMBERING_PATTERN,
    FONT_SIZE_TOLERANCE_PT,
    KEYWORD_SECTION_MARKERS,
    MARGIN_TOLERANCE_CM,
    MAX_KEYWORD_COUNT,
    MAX_UNJUSTIFIED_PARAGRAPH_RATIO,
    MIN_KEYWORD_COUNT,
    MIN_WORDS_FOR_ABSTRACT_CHECK,
    REQUIRED_FONT_SIZE_PT,
    REQUIRED_MARGIN_CM,
    TABLE_CAPTION_PREFIXES,
    TABLE_NUMBERING_PATTERN,
)
from src.infrastructure.adapters.document.eumic_violation_factory import EumicViolationFactory


class DocxEumicAdapter(DocumentFormatInspectionPort):
    """Inspects a .docx document for EUMIC editorial standard violations."""

    def inspect(self, docx_path: str, word_count: int) -> list[EumicViolationDTO]:
        document = Document(docx_path)
        violations: list[EumicViolationDTO] = []
        violations.extend(self._verify_format(document=document))
        violations.extend(self._verify_figures(document=document))
        violations.extend(self._verify_tables(document=document))
        violations.extend(self._verify_formulas(document=document))
        violations.extend(self._verify_abstract_keywords(document=document, word_count=word_count))
        return violations

    def _verify_format(self, document) -> list[EumicViolationDTO]:
        violations: list[EumicViolationDTO] = []
        violations.extend(self._check_margins(document=document))
        violations.extend(self._check_fonts(document=document))
        violations.extend(self._check_text_alignment(document=document))
        return violations

    def _verify_figures(self, document) -> list[EumicViolationDTO]:
        image_count = self._count_image_relationships(document=document)
        if image_count == 0:
            return []
        figure_captions = self._collect_paragraphs_starting_with(
            document=document, prefixes=FIGURE_CAPTION_PREFIXES
        )
        violations: list[EumicViolationDTO] = []
        violations.extend(
            self._check_figure_caption_count(
                image_count=image_count, caption_count=len(figure_captions)
            )
        )
        has_numbering_violations = len(
            figure_captions
        ) > 1 and self._has_sequential_numbering_violations(
            captions=figure_captions, pattern=FIGURE_NUMBERING_PATTERN
        )
        if has_numbering_violations:
            violations.append(EumicViolationFactory.figures_inconsistent_numbering())
        return violations

    def _verify_tables(self, document) -> list[EumicViolationDTO]:
        tables = document.tables
        if not tables:
            return []
        table_titles = self._collect_paragraphs_starting_with(
            document=document, prefixes=TABLE_CAPTION_PREFIXES
        )
        violations: list[EumicViolationDTO] = []
        violations.extend(
            self._check_table_title_count(table_count=len(tables), title_count=len(table_titles))
        )
        has_numbering_violations = len(
            table_titles
        ) > 1 and self._has_sequential_numbering_violations(
            captions=table_titles, pattern=TABLE_NUMBERING_PATTERN
        )
        if has_numbering_violations:
            violations.append(EumicViolationFactory.tables_inconsistent_numbering())
        return violations

    def _verify_formulas(self, document) -> list[EumicViolationDTO]:
        formula_paragraphs = self._collect_formula_paragraphs(document=document)
        if not formula_paragraphs:
            return []
        return self._check_formula_alignment(formula_paragraphs=formula_paragraphs)

    def _verify_abstract_keywords(self, document, word_count: int) -> list[EumicViolationDTO]:
        if word_count < MIN_WORDS_FOR_ABSTRACT_CHECK:
            return []
        violations: list[EumicViolationDTO] = []
        has_abstract, abstract_word_count = self._find_abstract_word_count(document=document)
        violations.extend(
            self._check_abstract(has_abstract=has_abstract, abstract_word_count=abstract_word_count)
        )
        has_keywords, keyword_count = self._find_keyword_count(document=document)
        violations.extend(
            self._check_keywords(has_keywords=has_keywords, keyword_count=keyword_count)
        )
        return violations

    def _check_margins(self, document) -> list[EumicViolationDTO]:
        sections = document.sections
        if not sections:
            return []
        section = sections[0]
        required_twips = Cm(REQUIRED_MARGIN_CM).twips
        tolerance_twips = Cm(MARGIN_TOLERANCE_CM).twips
        margins = [
            ("superior", section.top_margin),
            ("inferior", section.bottom_margin),
            ("izquierdo", section.left_margin),
            ("derecho", section.right_margin),
        ]
        return [
            EumicViolationFactory.margin_non_compliant(
                margin_name=margin_name,
                actual_cm=margin_value.cm,
                required_cm=REQUIRED_MARGIN_CM,
            )
            for margin_name, margin_value in margins
            if abs(margin_value.twips - required_twips) > tolerance_twips
        ]

    def _check_fonts(self, document) -> list[EumicViolationDTO]:
        all_runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
        fonts_used = {run.font.name for run in all_runs if run.font.name}
        sizes_used = {run.font.size for run in all_runs if run.font.size}
        violations: list[EumicViolationDTO] = []
        non_standard_fonts = fonts_used - set(AllowedFont)
        if non_standard_fonts:
            violations.append(
                EumicViolationFactory.non_standard_fonts(detected_fonts=non_standard_fonts)
            )
        non_standard_sizes = [
            size
            for size in sizes_used
            if abs(size.pt - REQUIRED_FONT_SIZE_PT) > FONT_SIZE_TOLERANCE_PT
        ]
        if non_standard_sizes:
            violations.append(
                EumicViolationFactory.variable_font_sizes(non_standard_sizes=non_standard_sizes)
            )
        return violations

    def _check_text_alignment(self, document) -> list[EumicViolationDTO]:
        non_justified = 0
        total_paragraphs = 0
        for paragraph in document.paragraphs:
            if not paragraph.text.strip():
                continue
            total_paragraphs += 1
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                non_justified += 1
        if (
            total_paragraphs > 0
            and non_justified / total_paragraphs > MAX_UNJUSTIFIED_PARAGRAPH_RATIO
        ):
            return [
                EumicViolationFactory.text_not_justified(
                    non_justified=non_justified, total=total_paragraphs
                )
            ]
        return []

    def _count_image_relationships(self, document) -> int:
        count = 0
        try:
            for relationship in document.part.rels.values():
                if "image" in relationship.target_ref:
                    count += 1
        except (KeyError, AttributeError):
            pass
        return count

    def _collect_paragraphs_starting_with(self, document, prefixes: tuple[str, ...]) -> list[str]:
        return [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip().lower().startswith(prefixes)
        ]

    def _check_figure_caption_count(
        self, image_count: int, caption_count: int
    ) -> list[EumicViolationDTO]:
        if caption_count >= image_count:
            return []
        return [EumicViolationFactory.figures_no_title(image_count=image_count)]

    def _has_sequential_numbering_violations(self, captions: list[str], pattern: str) -> bool:
        expected_number = 1
        for caption in captions:
            match = search(pattern, caption.lower())
            if not match:
                continue
            if int(match.group(1)) != expected_number:
                return True
            expected_number += 1
        return False

    def _check_table_title_count(
        self, table_count: int, title_count: int
    ) -> list[EumicViolationDTO]:
        if title_count >= table_count:
            return []
        return [
            EumicViolationFactory.tables_no_title(table_count=table_count, title_count=title_count)
        ]

    def _collect_formula_paragraphs(self, document) -> list:
        return [
            paragraph
            for paragraph in document.paragraphs
            if self._paragraph_contains_formula(paragraph=paragraph)
        ]

    def _paragraph_contains_formula(self, paragraph) -> bool:
        return any(self._run_contains_omath(run=run) for run in paragraph.runs)

    def _run_contains_omath(self, run) -> bool:
        try:
            xml_string = run._element.xml
            if isinstance(xml_string, bytes):
                xml_string = xml_string.decode("utf-8")
            return any(marker in xml_string for marker in FormulaXmlMarker)
        except AttributeError:
            return False

    def _check_formula_alignment(self, formula_paragraphs: list) -> list[EumicViolationDTO]:
        total = len(formula_paragraphs)
        unaligned_count = sum(
            1
            for paragraph in formula_paragraphs
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER
        )
        if unaligned_count == 0:
            return []
        return [EumicViolationFactory.formulas_not_centered(unaligned=unaligned_count, total=total)]

    def _find_abstract_word_count(self, document) -> tuple[bool, int]:
        paragraphs = document.paragraphs
        for paragraph_index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not any(keyword in text.lower() for keyword in ABSTRACT_SECTION_KEYWORDS):
                continue
            end_index = min(paragraph_index + ABSTRACT_PARAGRAPH_LOOKAHEAD, len(paragraphs))
            abstract_text = " ".join(
                paragraphs[index].text for index in range(paragraph_index, end_index)
            )
            return True, len(abstract_text.split())
        return False, 0

    def _check_abstract(
        self, has_abstract: bool, abstract_word_count: int
    ) -> list[EumicViolationDTO]:
        if not has_abstract:
            return [EumicViolationFactory.missing_abstract()]
        if (
            abstract_word_count < ABSTRACT_MIN_WORD_COUNT
            or abstract_word_count > ABSTRACT_MAX_WORD_COUNT
        ):
            return [
                EumicViolationFactory.abstract_length_out_of_range(
                    word_count=abstract_word_count,
                    min_words=ABSTRACT_MIN_WORD_COUNT,
                    max_words=ABSTRACT_MAX_WORD_COUNT,
                )
            ]
        return []

    def _find_keyword_count(self, document) -> tuple[bool, int]:
        for paragraph in document.paragraphs:
            text = paragraph.text.lower()
            if not any(marker in text for marker in KEYWORD_SECTION_MARKERS):
                continue
            keyword_text = text.split(":", 1)[-1] if ":" in text else text
            count = len([token for token in split(r"[,;]", keyword_text) if token.strip()])
            return True, count
        return False, 0

    def _check_keywords(self, has_keywords: bool, keyword_count: int) -> list[EumicViolationDTO]:
        if not has_keywords:
            return [EumicViolationFactory.missing_keywords()]
        if keyword_count < MIN_KEYWORD_COUNT or keyword_count > MAX_KEYWORD_COUNT:
            return [
                EumicViolationFactory.incorrect_keyword_count(
                    count=keyword_count,
                    min_count=MIN_KEYWORD_COUNT,
                    max_count=MAX_KEYWORD_COUNT,
                )
            ]
        return []
