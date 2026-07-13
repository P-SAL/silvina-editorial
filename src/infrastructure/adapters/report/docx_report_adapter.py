import re
from collections import defaultdict
from datetime import datetime

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.domain.dtos.report_input_dto import ReportInputDTO
from src.domain.enums.publication_verdict import PublicationVerdict
from src.domain.enums.recommendation_priority import RecommendationPriority
from src.domain.exceptions.report_errors import ReportExportUnavailable
from src.domain.report.report_export_port import ReportExportPort
from src.infrastructure.adapters.report.docx_report_settings import DocxReportSettings

_MARKDOWN_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


class DocxReportAdapter(ReportExportPort):
    """Export analysis reports to professionally formatted Word (.docx) documents."""

    def __init__(
        self,
        settings: DocxReportSettings,
        logo_path: str | None = None,
    ) -> None:
        if not DOCX_AVAILABLE:
            raise ReportExportUnavailable()
        self._logo_path = logo_path
        self._settings = settings

    def _color_for_score(self, score: float) -> tuple[int, int, int]:
        if score >= self._settings.score_high_threshold:
            return self._settings.publishable_color_rgb
        if score >= self._settings.score_medium_threshold:
            return self._settings.warning_color_rgb
        return self._settings.reject_color_rgb

    def _add_markdown_paragraph(self, doc, text: str) -> None:
        """Add a paragraph rendering **bold** Markdown segments as bold runs."""
        paragraph = doc.add_paragraph()
        position = 0
        for match in _MARKDOWN_BOLD_PATTERN.finditer(text):
            if match.start() > position:
                paragraph.add_run(text[position : match.start()])
            paragraph.add_run(match.group(1)).bold = True
            position = match.end()
        if position < len(text):
            paragraph.add_run(text[position:])

    def _format_match_rate(self, citations) -> str:
        if citations.total_citations == 0:
            return "N/A"
        rate = citations.matched_count / citations.total_citations * 100
        return f"{rate:.1f}%"

    def export(self, report_input: ReportInputDTO, output_path: str) -> bool:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = self._settings.font_name
        style.font.size = Pt(self._settings.base_font_size_pt)
        style.paragraph_format.line_spacing = self._settings.line_spacing
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

        self._add_header_logo(doc=doc)
        self._add_page_numbers(doc=doc)
        self._add_title_page(doc=doc, report_input=report_input)
        self._add_executive_summary(doc=doc, report_input=report_input)
        self._add_document_info(doc=doc, report_input=report_input)
        self._add_classification(doc=doc, report_input=report_input)
        self._add_quality_analysis(doc=doc, report_input=report_input)
        self._add_editorial_suitability(doc=doc, report_input=report_input)
        self._add_grammar_analysis(doc=doc, report_input=report_input)
        self._add_apa_validation(doc=doc, report_input=report_input)
        self._add_structure_validation(doc=doc, report_input=report_input)
        self._add_citations_analysis(doc=doc, report_input=report_input)
        self._add_recommendations(doc=doc, report_input=report_input)
        self._add_footer(doc=doc)

        doc.save(output_path)
        return True

    def _add_title_page(self, doc, report_input: ReportInputDTO) -> None:
        title = doc.add_heading("INFORME DE ANÁLISIS EDITORIAL", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(self._settings.title_font_size_pt)
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)
            run.font.bold = True
            run.font.underline = True

        doc.add_paragraph()

        doc_title = report_input.document_content.title or report_input.filename
        title_paragraph = doc.add_paragraph()
        title_paragraph.add_run(doc_title).bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.runs[0].font.size = Pt(self._settings.base_font_size_pt)

    def _add_header_logo(self, doc) -> None:
        if self._logo_path is None:
            return
        try:
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].clear()

            table = header.add_table(
                rows=1,
                cols=2,
                width=Inches(self._settings.header_table_width_inches),
            )

            left_cell = table.rows[0].cells[0]
            left_cell.width = Inches(self._settings.header_left_cell_width_inches)
            left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            left_paragraph = left_cell.paragraphs[0]
            left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            app_label_run = left_paragraph.add_run(
                f"Generado por {self._settings.app_name} v{self._settings.app_version}\n"
            )
            app_label_run.italic = True
            app_label_run.font.size = Pt(self._settings.metadata_font_size_pt)
            app_label_run.font.color.rgb = RGBColor(*self._settings.neutral_color_rgb)

            date_run = left_paragraph.add_run(datetime.now().strftime("%d de %B de %Y"))
            date_run.font.size = Pt(self._settings.metadata_font_size_pt)
            date_run.font.color.rgb = RGBColor(*self._settings.neutral_color_rgb)

            right_cell = table.rows[0].cells[1]
            right_cell.width = Inches(self._settings.header_right_cell_width_inches)
            right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            right_paragraph = right_cell.paragraphs[0]
            right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            logo_run = right_paragraph.add_run()
            logo_run.add_picture(self._logo_path, width=Inches(self._settings.logo_width_inches))

            table_element = table._element
            table_properties = table_element.tblPr
            if table_properties is None:
                table_properties = OxmlElement("w:tblPr")
                table_element.insert(0, table_properties)

            table_borders = OxmlElement("w:tblBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                table_borders.append(border)
            table_properties.append(table_borders)

        except (FileNotFoundError, OSError):
            pass  # Logo is decorative; skip if file is missing or unreadable

    def _add_page_numbers(self, doc) -> None:
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].clear()

        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        page_number_run = paragraph.add_run()

        page_field_begin = OxmlElement("w:fldChar")
        page_field_begin.set(qn("w:fldCharType"), "begin")
        page_number_run._element.append(page_field_begin)

        page_instruction = OxmlElement("w:instrText")
        page_instruction.set(qn("xml:space"), "preserve")
        page_instruction.text = "PAGE"
        page_number_run._element.append(page_instruction)

        page_field_end = OxmlElement("w:fldChar")
        page_field_end.set(qn("w:fldCharType"), "end")
        page_number_run._element.append(page_field_end)

        separator_run = paragraph.add_run(" de ")
        separator_run.font.size = Pt(self._settings.page_number_font_size_pt)

        num_pages_run = paragraph.add_run()

        num_pages_field_begin = OxmlElement("w:fldChar")
        num_pages_field_begin.set(qn("w:fldCharType"), "begin")
        num_pages_run._element.append(num_pages_field_begin)

        num_pages_instruction = OxmlElement("w:instrText")
        num_pages_instruction.set(qn("xml:space"), "preserve")
        num_pages_instruction.text = "NUMPAGES"
        num_pages_run._element.append(num_pages_instruction)

        num_pages_field_end = OxmlElement("w:fldChar")
        num_pages_field_end.set(qn("w:fldCharType"), "end")
        num_pages_run._element.append(num_pages_field_end)

        for text_run in paragraph.runs:
            text_run.font.size = Pt(self._settings.page_number_font_size_pt)
            text_run.font.color.rgb = RGBColor(*self._settings.neutral_color_rgb)

    def _add_executive_summary(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("RESUMEN EJECUTIVO", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        can_publish = report_input.is_publishable
        reason = report_input.publishability_reason

        decision_paragraph = doc.add_paragraph()
        decision_run = decision_paragraph.add_run(
            "✅ APTO PARA PUBLICACIÓN" if can_publish else "⚠️ REQUIERE REVISIÓN"
        )
        decision_run.bold = True
        decision_run.font.size = Pt(self._settings.decision_font_size_pt)
        decision_run.font.color.rgb = (
            RGBColor(*self._settings.publishable_color_rgb)
            if can_publish
            else RGBColor(*self._settings.reject_color_rgb)
        )

        doc.add_paragraph(reason)
        doc.add_paragraph()

        table = doc.add_table(rows=6, cols=2)
        table.style = self._settings.table_style

        header_cells = table.rows[0].cells
        header_cells[0].text = "Métrica"
        header_cells[1].text = "Valor"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        apa_count = len(report_input.apa_validation.violations)

        table.rows[1].cells[0].text = "Calidad General"
        table.rows[1].cells[1].text = f"{report_input.quality.overall_score:.1f}/10"

        table.rows[2].cells[0].text = "Gramática y Ortografía"
        table.rows[2].cells[1].text = f"{report_input.grammar.score:.1f}/10"

        table.rows[3].cells[0].text = "Estructura"
        table.rows[3].cells[1].text = (
            "✓ Válida" if report_input.structure.is_valid else "✗ Incompleta"
        )

        table.rows[4].cells[0].text = "Errores APA 7"
        table.rows[4].cells[1].text = f"{apa_count} detectados" if apa_count > 0 else "Sin errores"

        table.rows[5].cells[0].text = "Tasa de Coincidencia"
        table.rows[5].cells[1].text = self._format_match_rate(citations=report_input.citations)

        doc.add_paragraph()

    def _add_document_info(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("📄 INFORMACIÓN DEL DOCUMENTO", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        doc_content = report_input.document_content
        estimated_pages = doc_content.word_count // self._settings.words_per_page

        paragraph = doc.add_paragraph()
        paragraph.add_run("Título: ").bold = True
        paragraph.add_run(doc_content.title or report_input.filename)

        paragraph = doc.add_paragraph()
        paragraph.add_run("Autor(es): ").bold = True
        paragraph.add_run(doc_content.authors or "No identificado")

        paragraph = doc.add_paragraph()
        paragraph.add_run("Total de palabras: ").bold = True
        paragraph.add_run(f"{doc_content.word_count:,}")

        paragraph = doc.add_paragraph()
        paragraph.add_run("Total de caracteres: ").bold = True
        paragraph.add_run(f"{doc_content.char_count:,}")

        paragraph = doc.add_paragraph()
        paragraph.add_run("Páginas estimadas: ").bold = True
        paragraph.add_run(str(estimated_pages))

    def _add_classification(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("🏷️ CLASIFICACIÓN DEL ARTÍCULO", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        classification = report_input.classification

        paragraph = doc.add_paragraph()
        paragraph.add_run("Categoría: ").bold = True
        paragraph.add_run(classification.article_type.value.upper())

        paragraph = doc.add_paragraph()
        paragraph.add_run("Confianza: ").bold = True
        confidence_text = (
            f"{classification.confidence:.1%}" if classification.confidence is not None else "—"
        )
        paragraph.add_run(confidence_text)

        if classification.reasoning:
            paragraph = doc.add_paragraph()
            paragraph.add_run("Razonamiento: ").bold = True
            paragraph.add_run(classification.reasoning)

    def _add_quality_analysis(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("⭐ ANÁLISIS DE CALIDAD SEMÁNTICA", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        quality = report_input.quality

        paragraph = doc.add_paragraph()
        paragraph.add_run("Puntuación General: ").bold = True
        score_run = paragraph.add_run(f"{quality.overall_score:.1f}/10")
        score_run.bold = True
        score_run.font.size = Pt(self._settings.score_font_size_pt)
        score_run.font.color.rgb = RGBColor(*self._color_for_score(score=quality.overall_score))

        if quality.dimension_scores:
            for dim_name, dim_data in quality.dimension_scores.items():
                doc.add_heading(dim_name.capitalize(), level=3)

                paragraph = doc.add_paragraph()
                paragraph.add_run("Puntuación: ").bold = True
                paragraph.add_run(f"{dim_data['score']:.1f}/10")

                if dim_data.get("feedback"):
                    self._add_markdown_paragraph(doc=doc, text=dim_data["feedback"])

    def _add_editorial_suitability(self, doc, report_input: ReportInputDTO) -> None:
        suitability = report_input.quality.editorial_suitability
        if suitability is None:
            return

        heading = doc.add_heading("🎯 PERTINENCIA EDITORIAL", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        doc.add_heading("Contribución", level=3)
        paragraph = doc.add_paragraph()
        paragraph.add_run("Veredicto: ").bold = True
        paragraph.add_run(suitability.contribution_verdict)

        if suitability.contribution_phrase:
            paragraph = doc.add_paragraph()
            paragraph.add_run("Aporte identificado: ").bold = True
            paragraph.add_run(suitability.contribution_phrase)

        paragraph = doc.add_paragraph()
        paragraph.add_run("Observación: ").bold = True
        paragraph.add_run(suitability.contribution_observation)

        doc.add_heading("Alineación con líneas de investigación", level=3)
        paragraph = doc.add_paragraph()
        paragraph.add_run("Veredicto: ").bold = True
        paragraph.add_run(suitability.alignment_verdict)

        paragraph = doc.add_paragraph()
        paragraph.add_run("Líneas relacionadas: ").bold = True
        paragraph.add_run(suitability.alignment_lines)

        paragraph = doc.add_paragraph()
        paragraph.add_run("Justificación: ").bold = True
        paragraph.add_run(suitability.alignment_justification)

    def _add_grammar_analysis(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("📝 GRAMÁTICA Y ORTOGRAFÍA", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        grammar = report_input.grammar

        paragraph = doc.add_paragraph()
        paragraph.add_run("Puntuación: ").bold = True
        paragraph.add_run(f"{grammar.score:.1f}/10")

        paragraph = doc.add_paragraph()
        paragraph.add_run("Estado: ").bold = True
        paragraph.add_run(grammar.feedback)

        if grammar.errors:
            doc.add_paragraph()
            doc.add_paragraph("Errores Detectados:").bold = True

            max_errors = self._settings.max_errors_displayed
            context_limit = self._settings.context_truncation_limit
            for err in grammar.errors[:max_errors]:
                doc.add_paragraph(err.message, style="List Number")

                context_text = (
                    err.context
                    if len(err.context) < context_limit
                    else err.context[:context_limit] + "..."
                )
                doc.add_paragraph(f'   Contexto: "{context_text}"', style="List Bullet 2")

                if err.replacements:
                    max_replacements = self._settings.max_replacements
                    doc.add_paragraph(
                        f"   Sugerencia: {', '.join(err.replacements[:max_replacements])}",
                        style="List Bullet 2",
                    )

    def _add_apa_validation(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("📖 VALIDACIÓN APA 7 (ESPAÑOL)", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        violations = report_input.apa_validation.violations

        if not violations:
            no_citations = report_input.citations.total_citations == 0
            message = (
                "ℹ️ Sin citaciones en texto detectadas — no se pudo validar formato APA 7"
                if no_citations
                else "✅ Sin errores de formato APA 7 detectados"
            )
            color = (
                self._settings.neutral_color_rgb
                if no_citations
                else self._settings.publishable_color_rgb
            )
            doc.add_paragraph().add_run(message).font.color.rgb = RGBColor(*color)
            doc.add_paragraph()
            return

        by_type = defaultdict(list)
        for violation in violations:
            by_type[violation.error_type].append(violation)

        max_errors = self._settings.max_errors_displayed
        for _, errors in by_type.items():
            doc.add_paragraph(f"CITACIÓN INCORRECTA ({len(errors)}):", style="Heading 3")

            for index, err in enumerate(errors[:max_errors], 1):
                doc.add_paragraph(f"{index}. Citación: {err.citation_text}")
                doc.add_paragraph(f'   Ubicación: "{err.location}"', style="List Bullet 2")
                doc.add_paragraph(f"   Problema: {err.explanation}", style="List Bullet 2")
                if err.correction:
                    doc.add_paragraph(f"   Corrección: {err.correction}", style="List Bullet 2")

    def _add_structure_validation(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("📋 VALIDACIÓN DE ESTRUCTURA (EUMIC)", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        structure = report_input.structure

        if structure.is_valid:
            doc.add_paragraph().add_run(
                "✓ Estructura válida según normas EUMIC"
            ).font.color.rgb = RGBColor(*self._settings.publishable_color_rgb)
            return

        doc.add_paragraph().add_run("✗ Estructura incompleta").font.color.rgb = RGBColor(
            *self._settings.reject_color_rgb
        )

        if not structure.missing_sections:
            return

        doc.add_paragraph()
        doc.add_paragraph("Secciones faltantes:").bold = True
        for section in structure.missing_sections:
            doc.add_paragraph(f"• {section}", style="List Bullet")

    def _add_citations_analysis(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("📚 ANÁLISIS DE CITAS Y REFERENCIAS", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        citations = report_input.citations

        table = doc.add_table(rows=4, cols=2)
        table.style = self._settings.table_style

        table.rows[0].cells[0].text = "Total de citas en texto"
        table.rows[0].cells[1].text = str(citations.total_citations)

        table.rows[1].cells[0].text = "Total de referencias bibliográficas"
        table.rows[1].cells[1].text = str(citations.total_references)

        table.rows[2].cells[0].text = "Citas con referencia"
        table.rows[2].cells[1].text = str(citations.matched_count)

        table.rows[3].cells[0].text = "Tasa de coincidencia"
        table.rows[3].cells[1].text = self._format_match_rate(citations=citations)

    def _add_recommendations(self, doc, report_input: ReportInputDTO) -> None:
        heading = doc.add_heading("💡 RECOMENDACIONES", 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._settings.heading_color_rgb)

        verdict = report_input.verdict
        verdict_colors = {
            PublicationVerdict.CRITICAL: self._settings.reject_color_rgb,
            PublicationVerdict.WARNING: self._settings.reject_color_rgb,
            PublicationVerdict.APPROVED: self._settings.publishable_color_rgb,
        }
        paragraph = doc.add_paragraph()
        paragraph.add_run(verdict.message).bold = True
        paragraph.runs[0].font.size = Pt(self._settings.recommendation_font_size_pt)
        paragraph.runs[0].font.color.rgb = RGBColor(*verdict_colors[verdict.verdict])

        if not report_input.recommendations:
            return

        priority_icons = {
            RecommendationPriority.HIGH: "🔴",
            RecommendationPriority.MEDIUM: "🟡",
            RecommendationPriority.LOW: "🟢",
        }
        doc.add_paragraph("Recomendaciones específicas:").bold = True
        for rec in report_input.recommendations:
            icon = priority_icons.get(rec.priority, "⚪")
            doc.add_paragraph(f"{icon} {rec.message}", style="List Bullet")

    def _add_footer(self, doc) -> None:
        doc.add_paragraph()
        doc.add_paragraph("─" * 80)

        footer_paragraph = doc.add_paragraph()
        footer_paragraph.add_run(
            f"Generado por {self._settings.app_name} v{self._settings.app_version} | "
        ).italic = True
        footer_paragraph.add_run(datetime.now().strftime("%d/%m/%Y %H:%M")).italic = True
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.runs[0].font.size = Pt(self._settings.metadata_font_size_pt)
        footer_paragraph.runs[0].font.color.rgb = RGBColor(*self._settings.neutral_color_rgb)
        footer_paragraph.runs[1].font.size = Pt(self._settings.metadata_font_size_pt)
        footer_paragraph.runs[1].font.color.rgb = RGBColor(*self._settings.neutral_color_rgb)
