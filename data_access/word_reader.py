"""
word_reader.py
Reads Word documents and extracts raw text content.
Part of Silvina Editorial Assistant v0.7
"""

from typing import List
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordReader:
    """Reads Word documents (.docx) and extracts text content."""
    
    def __init__(self):
        """Initialize the Word reader."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required to read Word documents. "
                "Install it with: pip install python-docx"
            )
    
    def read_word_document(self, file_path: str) -> List[str]:
        """
        Read a Word document and return list of paragraphs.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            List of paragraph texts (strings)
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a .docx file
        """
        # Validate file exists
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file extension
        if path.suffix.lower() != '.docx':
            raise ValueError(f"File must be a .docx Word document: {file_path}")
        
        try:
            # Open and read document
            doc = Document(file_path)
            
            # Extract all paragraph texts
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
            
            return paragraphs
            
        except Exception as e:
            raise RuntimeError(f"Error reading Word document: {e}")
    
    def read_document_with_styles(self, file_path: str) -> List[dict]:
        """
        Read document and include paragraph style information.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            List of dictionaries with 'text' and 'style' keys
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            paragraphs_with_style = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs_with_style.append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            return paragraphs_with_style
            
        except Exception as e:
            raise RuntimeError(f"Error reading Word document with styles: {e}")
    
    def get_document_properties(self, file_path: str) -> dict:
        """
        Extract document metadata/properties.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Dictionary with document properties
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            properties = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'paragraph_count': len(doc.paragraphs),
            }
            
            return properties
            
        except Exception as e:
            raise RuntimeError(f"Error reading document properties: {e}")


# Convenience function for quick reading
def read_word_file(file_path: str) -> List[str]:
    """
    Quick function to read a Word document.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        List of paragraph texts
    """
    reader = WordReader()
    return reader.read_word_document(file_path)