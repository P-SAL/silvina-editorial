"""
Silvina Editorial Assistant v0.6 - REDESIGNED (Self-Contained)
Citation Integrity + IMRyD Validation + Deterministic Classification

Author: Pablo Salonio
Repository: https://github.com/P-SAL/silvina-editorial
"""

from datetime import datetime
import re
import win32com.client
import pythoncom
import time
import os
from difflib import SequenceMatcher
from dataclasses import dataclass
from typing import List, Optional, Dict
from enum import Enum
import requests
import json
from tqdm import tqdm
from docx import Document as DocxDocument 
from docx.shared import RGBColor, Pt, Inches  
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement  
from docx.oxml.ns import qn  

# ============================================================
# CONFIGURATION
# ============================================================
OLLAMA_MODEL = 'llama3-gradient:8b-instruct-1048k-q5_K_M'  # Change here to switch models
OLLAMA_URL = 'http://localhost:11434/api/generate'


# ============================================================
# LLM INTEGRATION (OLLAMA)
# ============================================================


def analyze_with_ollama(content: str, article_type: str, tier1_report: str) -> str:
    """Send content to Ollama with streaming progress."""
    
    prompt = f"""Eres un revisor editorial. Analiza este artículo de {article_type}.

TEXTO:
{content[:6000]}

Evalúa brevemente (máximo 200 palabras):
1. Claridad del argumento
2. Tono académico
3. Coherencia

Responde en español, sé directo. NO incluyas recomendaciones ni sugerencias."""
    
    try:
        print("⏳ Analizando con Ollama...")
        
        response = requests.post(
            OLLAMA_URL,  # ← Use constant
            json={
                'model': OLLAMA_MODEL,  # ← Use constant
                'prompt': prompt,
                'stream': True,
                'options': {
                    'temperature': 0.3,
                    'num_predict': 500
                }
            },
            stream=True,
            timeout=300
        )
        
        full_response = []
        
        with tqdm(desc="Generando análisis", unit=" tokens", bar_format='{l_bar}{bar}| {n_fmt} tokens') as pbar:
            for line in response.iter_lines():
                if line:
                    chunk = json.loads(line)
                    if 'response' in chunk:
                        token = chunk['response']
                        full_response.append(token)
                        pbar.update(1)
                    if chunk.get('done', False):
                        break
        
        print("✅ Análisis completado")
        return ''.join(full_response)
    
    except requests.exceptions.ConnectionError:
        return "❌ No conecta a Ollama (¿está corriendo 'ollama serve'?)"
    except Exception as e:
        return f"❌ Error: {str(e)}"



# ============================================================
# ENUMS AND DATA CLASSES
# ============================================================

class ArticleType(Enum):
    """Article types according to EUMIC guidelines."""
    CIENTIFICA = "Científica"
    DIVULGACION = "Divulgación"
    INDETERMINADO = "Indeterminado"


class SeverityLevel(Enum):
    """Issue severity levels for reporting."""
    CRITICO = "🔴 CRÍTICO"
    ADVERTENCIA = "🟡 ADVERTENCIA"
    INFORMATIVO = "🔵 INFORMATIVO"


@dataclass
class Citation:
    """Represents one in-text citation."""
    authors: List[str]
    year: str
    page: Optional[str]
    paragraph_index: int
    citation_type: str  # "narrativa" or "parentética"
    raw_text: str
    
    @property
    def key(self) -> str:
        """Generate matching key (handles abbreviations and last names)."""
        first_author = self.authors[0].replace(" et al.", "").strip()
        
        # Check if it's an abbreviation with dash: "CIA-" → "cia"
        if re.match(r'^[A-Z]{2,}[-–—]?$', first_author):
            key_word = first_author.rstrip('-–—').lower()
        else:
            # Use last word for organizations or just the name
            words = first_author.split()
            skip_words = {'de', 'del', 'la', 'el', 'los', 'las', 'y', 'Diario'}
            significant_words = [w for w in words if w not in skip_words]
            key_word = (significant_words[-1] if significant_words else words[-1]).lower()
        
        return f"{key_word}_{self.year}"
    
    def __repr__(self):
        authors_text = " y ".join(self.authors)
        page_info = f", p. {self.page}" if self.page else ""
        return f"{'📖' if self.citation_type == 'narrativa' else '📎'} {authors_text} ({self.year}{page_info}) [¶{self.paragraph_index}]"


@dataclass
class Section:
    """Represents a document section (e.g., Introducción, Métodos)."""
    name: str
    paragraph_index: int
    word_count: int
    expected_order: int
    
    def __repr__(self):
        return f"{self.name} @ ¶{self.paragraph_index} ({self.word_count} palabras)"


# ============================================================
# REFERENCE CLASS (From v0.5 - Included for self-containment)
# ============================================================

class Reference:
    """Represents a single bibliographic reference (v0.5 proven class)."""
    
    def __init__(self, text):
        self.text = text
    
    def validate_author(self):
        """Check if reference has valid APA 7 Spanish author format."""
        # Remove leading bullets/dashes
        clean_text = re.sub(r'^[-–—•]+\s*', '', self.text)
        
        # Personal author patterns
        personal = r'[A-ZÁ-ÚÑ][a-zá-úñ]+(?:-[A-ZÁ-ÚÑ][a-zá-úñ]+)?,\s+[A-Z]\.'
        et_al = r'et\s+al\.'
        
        # Organizational patterns (broader)
        org1 = r'^[A-Z][A-Za-záéíóúñ\s&,\-]{10,}\.\s'
        org2 = r'^[A-Z][A-Za-záéíóúñ\s&,\-]{10,}\s+\(\d{4}'
        org3 = r'^(Ministerio|Centro|Instituto|Diario|La\s+Nación|Central)'  # Common Spanish orgs
        
        has_personal = bool(re.search(personal, clean_text))
        has_et_al = bool(re.search(et_al, clean_text, re.IGNORECASE))
        has_org1 = bool(re.search(org1, clean_text))
        has_org2 = bool(re.search(org2, clean_text))
        has_org3 = bool(re.search(org3, clean_text, re.IGNORECASE))
        
        # Accept if any organizational pattern OR personal author
        if (has_org1 or has_org2 or has_org3) and not has_personal:
            return True
        
        return has_personal or has_et_al
    
            
    def validate_year(self):
        """Check if reference has valid year format (flexible patterns)."""
        # Pattern 1: Standard (2020) or (2020a)
        pattern1 = r'\((\d{4}[a-z]?)\)'
        # Pattern 2: Range (1983-2003)
        pattern2 = r'\((\d{4})-\d{4}\)'
        # Pattern 3: Date format (2004, diciembre 15)
        pattern3 = r'\((\d{4}),\s+\w+'
        
        for pattern in [pattern1, pattern2, pattern3]:
            match = re.search(pattern, self.text)
            if match:
                year = match.group(1)[:4]  # Take first 4 digits
                return True, year
        
        return False, None
    
        
    def validar_conjuncion_espanola(self):
        """Verifica uso de 'y' en vez de '&' para referencias en español APA 7."""
        patron_ampersand = r'[A-Z]\.(?:,)?\s+&\s+[A-Z]'
        
        if re.search(patron_ampersand, self.text):
            return False, "Uso incorrecto de '&' (debe ser 'y' en español APA 7)"
        
        return True, None
    
    def tiene_doi_o_url(self):
        """Verifica presencia de DOI o URL."""
        tiene_doi = bool(re.search(r'https?://doi\.org/[\w\.\-/]+', self.text, re.IGNORECASE))
        tiene_url = bool(re.search(r'https?://[^\s]+', self.text))
        formato_antiguo = bool(re.search(r'Recuperado\s+de\s+https?://', self.text, re.IGNORECASE))
        
        return {
            'tiene_doi': tiene_doi,
            'tiene_url': tiene_url,
            'formato_antiguo': formato_antiguo
        }
    
    def is_valid(self):
        """Check if reference meets all APA 7 Spanish requirements."""
        has_author = self.validate_author()
        has_year, _ = self.validate_year()
        conjuncion_valida, _ = self.validar_conjuncion_espanola()
        
        return has_author and has_year and conjuncion_valida
    
    def get_validation_report(self):
        """Return detailed validation results."""
        has_author = self.validate_author()
        has_year, year = self.validate_year()
        conjuncion_valida, error_conjuncion = self.validar_conjuncion_espanola()
        doi_url_info = self.tiene_doi_o_url()
        
        return {
            'text': self.text[:80] + '...' if len(self.text) > 80 else self.text,
            'valid_author': has_author,
            'valid_year': has_year,
            'valid_conjuncion': conjuncion_valida,
            'error_conjuncion': error_conjuncion,
            'doi_url_info': doi_url_info,
            'year': year,
            'is_valid': has_author and has_year and conjuncion_valida
        }


# ============================================================
# CITATION EXTRACTOR
# ============================================================

class CitationExtractor:
    """Extracts APA citations from Spanish text."""
    
    # Consolidated patterns
    AUTHOR_PATTERN = r'[A-ZÁÉÍÓÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\-]+'
    YEAR_PATTERN = r'\d{4}[a-z]?'
    PAGE_PATTERN = r'(?:pp?\.|párr\.)\s*([\d\-]+)'
    
    def __init__(self):
        # Parenthetical: (García, 2020) or (Ministerio, 2020a) or (CIA, 1985)
        self.pattern_parenthetical = re.compile(
            rf'\(([A-ZÁ-ÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\-]+(?:\s+et\s+al\.)?(?:\s+y\s+[A-ZÁ-ÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\-]+)*),?\s+'
            rf'({self.YEAR_PATTERN})(?:,\s*{self.PAGE_PATTERN})?\)',
            re.IGNORECASE
        )
        
        # Narrative: García (2020) or Ministerio (2020a)
        self.pattern_narrative = re.compile(
            rf'([A-ZÁ-ÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\-]+(?:\s+et\s+al\.)?(?:\s+y\s+[A-ZÁ-ÚÑ][a-záéíóúñA-ZÁÉÍÓÚÑ\-]+)*)\s+'
            rf'\(({self.YEAR_PATTERN})(?:,\s*{self.PAGE_PATTERN})?\)',
            re.IGNORECASE
        )
    
        
    def extract_from_text(self, text: str, para_index: int) -> List[Citation]:
        """Extract all citations from one paragraph."""
        citations = []
        
        # Parenthetical citations
        for match in self.pattern_parenthetical.finditer(text):
            authors_raw = match.group(1)
            year = match.group(2)
            page = match.group(3) if match.lastindex >= 3 else None
            
            citations.append(Citation(
                authors=self._parse_authors(authors_raw),
                year=year,
                page=page,
                paragraph_index=para_index,
                citation_type="parentética",
                raw_text=match.group(0)
            ))
        
        # Narrative citations
        for match in self.pattern_narrative.finditer(text):
            authors_raw = match.group(1)
            year = match.group(2)
            page = match.group(3) if match.lastindex >= 3 else None
            
            citations.append(Citation(
                authors=self._parse_authors(authors_raw),
                year=year,
                page=page,
                paragraph_index=para_index,
                citation_type="narrativa",
                raw_text=match.group(0)
            ))
        
        return citations
    
    @staticmethod
    def _parse_authors(authors_text: str) -> List[str]:
        """Parse author string into list."""
        if "et al." in authors_text:
            first_author = authors_text.split("et al.")[0].strip()
            return [f"{first_author} et al."]
        
        if " y " in authors_text:
            return [a.strip() for a in authors_text.split(" y ")]
        
        return [authors_text.strip()]


# ============================================================
# STRUCTURE VALIDATOR
# ============================================================

class StructureValidator:
    """Validates IMRyD structure in scientific articles."""
    
    REQUIRED_SECTIONS = {
        "introducción": {"order": 1, "min_words": 300, "aliases": ["introduccion", "marco teórico", "marco teorico"]},
        "métodos": {"order": 2, "min_words": 200, "aliases": ["metodos", "metodología", "metodologia", "método"]},
        "resultados": {"order": 3, "min_words": 300, "aliases": ["resultados y análisis", "resultados y analisis"]},
        "discusión": {"order": 4, "min_words": 300, "aliases": ["discusion"]},
        "conclusiones": {"order": 5, "min_words": 150, "aliases": ["conclusión", "conclusion"]},
    }
    
    def extract_sections(self, paragraphs: List[str]) -> List[Section]:
        """Extract IMRyD sections from document."""
        sections = []
        current_section = None
        current_start = -1
        current_words = 0
        
        for i, para in enumerate(paragraphs):
            para_clean = para.strip().lower()
            
            # Skip empty paragraphs
            if not para_clean:
                continue
            
            # Check if paragraph is a section header
            section_found = None
            for section_name, section_info in self.REQUIRED_SECTIONS.items():
                if para_clean == section_name or para_clean in section_info["aliases"]:
                    section_found = section_name
                    break
            
            if section_found:
                # Save previous section
                if current_section:
                    sections.append(Section(
                        name=current_section,
                        paragraph_index=current_start,
                        word_count=current_words,
                        expected_order=self.REQUIRED_SECTIONS[current_section]["order"]
                    ))
                
                # Start new section
                current_section = section_found
                current_start = i
                current_words = 0
            else:
                # Accumulate words in current section
                if current_section:
                    current_words += len(para.split())
        
        # Save last section
        if current_section:
            sections.append(Section(
                name=current_section,
                paragraph_index=current_start,
                word_count=current_words,
                expected_order=self.REQUIRED_SECTIONS[current_section]["order"]
            ))
        
        return sections
    
    def validate(self, sections: List[Section]) -> Dict:
        """Validate structure and return issues."""
        issues = {
            "missing_sections": [],
            "out_of_order": [],
            "too_short": []
        }
        
        # Check for missing sections
        found_names = {s.name for s in sections}
        for required in self.REQUIRED_SECTIONS.keys():
            if required not in found_names:
                issues["missing_sections"].append(required)
        
        # Check order
        for i in range(len(sections) - 1):
            if sections[i].expected_order > sections[i + 1].expected_order:
                issues["out_of_order"].append((sections[i].name, sections[i + 1].name))
        
        # Check minimum length
        for section in sections:
            min_words = self.REQUIRED_SECTIONS[section.name]["min_words"]
            if section.word_count < min_words:
                issues["too_short"].append({
                    "section": section.name,
                    "current": section.word_count,
                    "minimum": min_words
                })
        
        return issues


# ============================================================
# ARTICLE CLASSIFIER
# ============================================================

class ArticleClassifier:
    """Deterministic article type classification."""
    
    # EUMIC thresholds
    CIENTIFICO_MIN_CHARS = 30000
    CIENTIFICO_MAX_CHARS = 50000
    DIVULGACION_TARGET_CHARS = 30000
    DIVULGACION_TOLERANCE = 5000
    
    MIN_CITATIONS_SCIENTIFIC = 5
    MIN_SECTIONS_SCIENTIFIC = 3
    MIN_BIBLIOGRAPHY_SCIENTIFIC = 1000  # characters
    
    @staticmethod
    def collect_metrics(doc_obj) -> Dict:
        """Collect all metrics needed for classification."""
        char_count = doc_obj.get_character_count()
        
        # Already extracted in load()
        citations_count = len(doc_obj.citations)
        sections_count = len(doc_obj.sections)
        bib_chars = len(doc_obj.text)
        
        return {
            'char_count': char_count,
            'citations_count': citations_count,
            'citation_density': (citations_count / (char_count / 1000)) if char_count > 0 else 0,
            'imryd_sections': sections_count,
            'section_names': [s.name for s in doc_obj.sections],
            'bibliography_chars': bib_chars,
            'bibliography_refs': len(doc_obj.references)
        }
    
    @classmethod
    def classify(cls, metrics: Dict) -> Dict:
        """
        Deterministic classification using EUMIC rules.
        NO LLM involvement - pure Python logic.
        """
        citations = metrics['citations_count']
        sections = metrics['imryd_sections']
        bib_chars = metrics['bibliography_chars']
        char_count = metrics['char_count']
        
        # === RULE 1: No citations = Divulgación (absolute rule) ===
        if citations == 0:
            return {
                'type': ArticleType.DIVULGACION,
                'confidence': 'alta',
                'score': 2,
                'reasons': {
                    'critical': ['Sin citas APA formales en texto'],
                    'positive': [] if bib_chars == 0 else ['Incluye bibliografía consultada'],
                    'length_valid': cls._check_divulgacion_length(char_count)
                }
            }
        
        # === RULE 2: Check scientific criteria ===
        meets_citations = citations >= cls.MIN_CITATIONS_SCIENTIFIC
        meets_structure = sections >= cls.MIN_SECTIONS_SCIENTIFIC
        meets_bibliography = bib_chars >= cls.MIN_BIBLIOGRAPHY_SCIENTIFIC
        
        scientific_score = (
            (4 if meets_citations else 0) +
            (3 if meets_structure else 0) +
            (3 if meets_bibliography else 0)
        )
        
        # === RULE 3: Classify based on score ===
        if scientific_score >= 8:
            # Scientific article
            length_valid = cls.CIENTIFICO_MIN_CHARS <= char_count <= cls.CIENTIFICO_MAX_CHARS
            
            return {
                'type': ArticleType.CIENTIFICA,
                'confidence': 'alta',
                'score': scientific_score,
                'reasons': {
                    'critical': [] if length_valid else [f'Longitud fuera de rango científico: {char_count:,} caracteres'],
                    'positive': [
                        f'{citations} citas APA detectadas',
                        f'{sections}/5 secciones IMRyD presentes',
                        f'Bibliografía extensa ({bib_chars} caracteres)'
                    ],
                    'length_valid': length_valid
                }
            }
        else:
            # Divulgación (failed scientific criteria)
            issues = []
            if not meets_citations:
                issues.append(f'Citas insuficientes: {citations} (mínimo: {cls.MIN_CITATIONS_SCIENTIFIC})')
            if not meets_structure:
                issues.append(f'Estructura IMRyD incompleta: {sections}/5 secciones')
            if not meets_bibliography:
                issues.append(f'Bibliografía menor a 1 página: {bib_chars} caracteres')
            
            return {
                'type': ArticleType.DIVULGACION,
                'confidence': 'media',
                'score': scientific_score,
                'reasons': {
                    'critical': issues,
                    'positive': [],
                    'length_valid': cls._check_divulgacion_length(char_count)
                }
            }
    
    @classmethod
    def _check_divulgacion_length(cls, char_count: int) -> bool:
        """Check if length is valid for divulgación."""
        return abs(char_count - cls.DIVULGACION_TARGET_CHARS) <= cls.DIVULGACION_TOLERANCE


# ============================================================
# CITATION MATCHER
# ============================================================

# ============================================================
# HYBRID ANALYSIS STRATEGY
# ============================================================

class HybridAnalysisStrategy:
    """Two-tier analysis: Full structural + Selective LLM quality."""
    
    FULL_LLM_THRESHOLD = 5000  # words
    
    @staticmethod
    def create_analysis_plan(word_count: int, sections: List[Section]) -> Dict:
        """Create analysis plan based on document length."""
        
        # TIER 1: Always full document (deterministic)
        tier1_plan = {
            'scope': 'FULL_DOCUMENT',
            'uses_llm': False
        }
        
        # TIER 2: Selective LLM analysis
        if word_count <= HybridAnalysisStrategy.FULL_LLM_THRESHOLD:
            tier2_plan = {
                'scope': 'FULL_DOCUMENT',
                'sections': 'ALL',
                'uses_llm': True
            }
        else:
            tier2_plan = {
                'scope': 'STRATEGIC_SAMPLING',
                'sections': HybridAnalysisStrategy._select_key_sections(sections),
                'uses_llm': True
            }
        
        return {
            'word_count': word_count,
            'tier1_structural': tier1_plan,
            'tier2_quality': tier2_plan
        }
    
    @staticmethod
    def _select_key_sections(sections: List[Section]) -> List[Dict]:
        """Select strategic sections for LLM analysis."""
        key_sections = []
        
        # Map section names to strategies
        section_strategies = {
            'introducción': 'FIRST_800_WORDS',
            'métodos': 'FIRST_300_WORDS',
            'resultados': 'FIRST_300_WORDS',
            'discusión': 'FIRST_300_WORDS',
            'conclusiones': 'FULL'
        }
        
        for section in sections:
            strategy = section_strategies.get(section.name.lower(), 'FIRST_300_WORDS')
            key_sections.append({
                'name': section.name,
                'paragraph_index': section.paragraph_index,
                'strategy': strategy
            })
        
        return key_sections



class CitationMatcher:
    """Matches in-text citations with reference list."""
    
    def __init__(self, citations: List[Citation], references: List[Reference]):
        self.citations = citations
        self.references = references
        
        # Build lookup keys WITH DEBUG
        self.citation_keys = {}
        for cit in citations:
            key = cit.key
            self.citation_keys[key] = cit
        
        self.reference_keys = {}
        for ref in references:
            key = self._ref_key(ref)
            self.reference_keys[key] = ref
        
          
    
    @staticmethod
    def _ref_key(reference: Reference) -> str:
        """Generate matching key with smart organizational handling."""
        # Remove bullets/dashes
        clean_text = re.sub(r'^[-–—•]+\s*', '', reference.text)
        
        # Try to match organizational pattern with year
        org_pattern = r'^([A-ZÁ-ÚÑ][A-Za-záéíóúñ\s&,\-]{5,}?)\s+\((\d{4}[a-z]?)'
        match = re.search(org_pattern, clean_text)
        
        if match:
            org_name = match.group(1).strip()
            year = match.group(2)  # ✅ Keep full year with letter

           
            # Extract last significant word (e.g., "Ministerio de Economía" → "economía")
            # or use abbreviated form if present
            
            # Check for abbreviation in parentheses
            abbrev_match = re.search(r'[-–—]\s*([A-Z]{2,})\s*[-–—]', org_name)
            if abbrev_match:
                # Use abbreviation: "Central Intelligence Agency -CIA-" → "cia"
                key_word = abbrev_match.group(1).lower()
            else:
                # Use last word: "Ministerio de Economía" → "economía"
                words = org_name.split()
                # Filter out common words
                skip_words = {'de', 'del', 'la', 'el', 'los', 'las', 'y'}
                significant_words = [w for w in words if w.lower() not in skip_words]
                key_word = significant_words[-1].lower() if significant_words else words[-1].lower()
            
            return f"{key_word}_{year}"
        
        # Fallback: personal author
        match = re.match(r'([A-ZÁ-ÚÑ][a-zá-úñ\-]+)', clean_text)
        if match:
            first_author = match.group(1).lower()
            _, year = reference.validate_year()
            return f"{first_author}_{year if year else 'unknown'}"
        
        return "unknown_unknown"
    
       
    def find_orphaned_citations(self) -> List[Citation]:
        """Citations without matching references."""
        orphaned = []
        for cit in self.citations:
            if cit.key not in self.reference_keys:
                orphaned.append(cit)
        return orphaned
    
    def find_orphaned_references(self) -> List[Reference]:
        """References never cited in text."""
        orphaned = []
        for ref in self.references:
            if self._ref_key(ref) not in self.citation_keys:
                orphaned.append(ref)
        return orphaned
    
    def generate_report(self, section_type: str) -> str:
        """Generate citation integrity report."""
        report = []
        report.append("\n" + "=" * 70)
        report.append("INTEGRIDAD DE CITAS Y REFERENCIAS")
        report.append("=" * 70)
        
        report.append(f"Citas en texto: {len(self.citations)}")
        report.append(f"Referencias bibliográficas: {len(self.references)}")
        report.append(f"Tipo de sección: {section_type.upper()}")
        
        # Orphaned citations (ALWAYS CRITICAL)
        orphaned_cits = self.find_orphaned_citations()
        if orphaned_cits:
            report.append(f"\n{SeverityLevel.CRITICO.value}: Citas Sin Referencia")
            report.append(f"Encontradas {len(orphaned_cits)} citas sin entrada bibliográfica:")
            for cit in orphaned_cits[:5]:
                report.append(f"  • {cit}")
            if len(orphaned_cits) > 5:
                report.append(f"  ... y {len(orphaned_cits) - 5} más")
        
        # Orphaned references (severity depends on section type)
        orphaned_refs = self.find_orphaned_references()
        if orphaned_refs:
            if section_type == "Referencias":
                severity = SeverityLevel.ADVERTENCIA
                msg = "En 'Referencias', se espera citar todas las entradas."
            else:
                severity = SeverityLevel.INFORMATIVO
                msg = "En 'Bibliografía', es aceptable incluir fuentes consultadas."
            
            report.append(f"\n{severity.value}: Referencias Sin Citar")
            report.append(msg)
            report.append(f"Encontradas {len(orphaned_refs)} referencias no citadas:")
            for ref in orphaned_refs[:5]:
                report.append(f"  • {ref.text[:60]}...")
            if len(orphaned_refs) > 5:
                report.append(f"  ... y {len(orphaned_refs) - 5} más")
        
        # Final verdict
        if not orphaned_cits and not orphaned_refs:
            report.append(f"\n✅ Sistema de citación íntegro")
        elif not orphaned_cits:
            report.append(f"\n✅ Todas las citas tienen referencia válida")
        
        return '\n'.join(report)


# ============================================================
# DOCUMENT CLASS (v0.5 + v0.6+)
# ============================================================

class Document:
    """Extended Document class with v0.6 capabilities."""
    
    def __init__(self, filepath):
        self.filepath = filepath
        self.word = None
        self.doc = None
        self.text = ""
        self.references = []
        self.section_type = "Referencias"
        
        # v0.6 additions
        self.citations = []
        self.sections = []
        self.article_classification = None

        # NEW: Add these two lines
        self.analysis_plan = None
        self.word_count = 0
    
    def load(self):
        """Load document and create analysis plan."""
        self._connect_to_word()
        self._extract_referencias()
        self._create_reference_objects()
        self._extract_citations()
        self._extract_structure()
        
        # NEW: Calculate word count and create analysis plan
        self.word_count = self._calculate_word_count()
        self.analysis_plan = HybridAnalysisStrategy.create_analysis_plan(
            self.word_count, 
            self.sections
        )
        
        print(f"\n📊 Plan de Análisis:")
        print(f"   Palabras: {self.word_count:,}")
        print(f"   Tier 1 (Estructural): Completo")
        print(f"   Tier 2 (Calidad): {self.analysis_plan['tier2_quality']['scope']}")
    
       
    def _connect_to_word(self):
        """Open Word document."""
        pythoncom.CoInitialize()
        
        try:
            self.word = win32com.client.Dispatch("Word.Application")
            self.word.Visible = False
            abs_path = os.path.abspath(self.filepath)
            self.doc = self.word.Documents.Open(abs_path)
            
            time.sleep(2.0)
            self.doc.Activate()
            time.sleep(1.0)
            
            print(f"✅ Documento cargado: {os.path.basename(self.filepath)}")
            
        except Exception as e:
            print(f"❌ Error de conexión: {e}")
            self.word = None
            self.doc = None
    
    def _extract_referencias(self):
        """Extract references section."""
        if not self.doc:
            return
        
        try:
            found_start = False
            referencias_paras = []
            
            for para in self.doc.Paragraphs:
                try:
                    para_text = para.Range.Text.strip()
                except:
                    continue
                
                if not found_start:
                    if "Bibliografía" in para_text or "Fuentes bibliográficas" in para_text:
                        self.section_type = "Bibliografía"
                        found_start = True
                        continue
                    elif "Referencias" in para_text and "bibliográficas" in para_text:
                        self.section_type = "Referencias"
                        found_start = True
                        continue
                
                if found_start and para_text and len(para_text) > 30:
                    referencias_paras.append(para_text)
            
            self.text = '\n'.join(referencias_paras)
            
            if len(referencias_paras) > 0:
                print(f"✅ Extraídas {len(referencias_paras)} referencias")
            else:
                print(f"⚠️  No se encontró sección de referencias/bibliografía")
            
        except Exception as e:
            print(f"❌ Error extrayendo referencias: {e}")
    
    def _create_reference_objects(self):
        """Create Reference objects from extracted text."""
        if not self.text:
            return
        
        paragraphs = self.text.split('\n')
        
        for para in paragraphs:
            para = para.strip()
            if len(para) < 30:
                continue
            
            # Check if paragraph has multiple references (separated by years)
            years = re.findall(r'\(\d{4}\)', para)
            
            if len(years) >= 2:
                # Try to split by period before capital letter
                split_pattern = r'\.(?=[A-Z][a-z]+,\s+[A-Z]\.)'
                parts = re.split(split_pattern, para, maxsplit=1)
                
                for part in parts:
                    part = part.strip()
                    if len(part) > 30:
                        if not part.endswith('.'):
                            part += '.'
                        self.references.append(Reference(part))
            else:
                self.references.append(Reference(para))
        
        if len(self.references) > 0:
            print(f"✅ Creados {len(self.references)} objetos Reference")
    
    
    def _extract_citations(self):
        """Extract all in-text citations with progress bar."""
        if not self.doc:
            return
        
        extractor = CitationExtractor()
        total_paras = self.doc.Paragraphs.Count
        
        print("⏳ Extrayendo citas del texto...")
        for i in tqdm(range(1, total_paras + 1), desc="Párrafos", unit="¶"):
            try:
                para = self.doc.Paragraphs(i)
                para_text = para.Range.Text.strip()
                if len(para_text) > 10:
                    citations = extractor.extract_from_text(para_text, i)
                    self.citations.extend(citations)
            except:
                continue
        
        if len(self.citations) > 0:
            print(f"✅ Extraídas {len(self.citations)} citas en texto")
        else:
            print(f"⚠️  No se encontraron citas APA en el texto")
    
    def _extract_structure(self):
        """Extract IMRyD structure with progress bar."""
        if not self.doc:
            return
        
        paragraphs = []
        total_paras = self.doc.Paragraphs.Count
        
        print("⏳ Analizando estructura IMRyD...")
        for i in tqdm(range(1, total_paras + 1), desc="Estructura", unit="¶"):
            try:
                para = self.doc.Paragraphs(i)
                para_text = para.Range.Text.strip()
                paragraphs.append(para_text)
            except:
                continue
        
        validator = StructureValidator()
        self.sections = validator.extract_sections(paragraphs)
        
        print(f"✅ Detectadas {len(self.sections)}/5 secciones IMRyD")

      
    def classify_article(self):
        """Classify article type deterministically."""
        metrics = ArticleClassifier.collect_metrics(self)
        self.article_classification = ArticleClassifier.classify(metrics)
        
        return self.article_classification
    
    def get_character_count(self):
        """Get accurate character count."""
        if not self.doc:
            return 0
        try:
            return self.doc.Characters.Count
        except:
            return 0
    
    def _calculate_word_count(self):
        """Calculate accurate word count."""
        if not self.doc:
            return 0
        try:
            return self.doc.Words.Count
        except:
            return 0
    
    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header."""
        text_lower = text.lower().strip()
        headers = ['introducción', 'métodos', 'resultados', 'discusión', 
                   'conclusiones', 'referencias', 'bibliografía']
        return text_lower in headers
    
    def get_llm_analysis_content(self) -> str:
        """Extract content for LLM analysis based on plan."""
        if not self.analysis_plan:
            return ""
        
        plan = self.analysis_plan['tier2_quality']
        
        
        if plan['scope'] == 'FULL_DOCUMENT':
            content = self._get_full_document_text()
            return content
        
        elif plan['scope'] == 'STRATEGIC_SAMPLING':
            content = self._build_strategic_excerpt(plan['sections'])
            return content
        
        return ""
    
    def _get_full_document_text(self) -> str:
        """Extract complete document text."""
        if not self.doc:
            return ""
        
        paragraphs = []
        try:
            total_paras = self.doc.Paragraphs.Count
                        
            for para in self.doc.Paragraphs:
                try:
                    text = para.Range.Text.strip()
                    if text:
                        paragraphs.append(text)
                except:
                    continue
            
            full_text = '\n\n'.join(paragraphs)
            return full_text
            
        except Exception as e:
            return ""

        
    def _build_strategic_excerpt(self, sections_plan: List[Dict]) -> str:
        """Build strategic excerpt for large documents."""
        content = []
        content.append("=== ANÁLISIS ESTRATÉGICO (DOCUMENTO LARGO) ===")
        content.append(f"Palabras totales: {self.word_count:,}")
        content.append(f"Validación estructural: COMPLETADA (Tier 1)\n")
        
        for section_info in sections_plan:
            section_text = self._extract_section_text(
                section_info['name'],
                section_info['paragraph_index'],
                section_info['strategy']
            )
            
            content.append(f"\n{'='*60}")
            content.append(f"{section_info['name'].upper()} [{section_info['strategy']}]")
            content.append('='*60)
            content.append(section_text)
        
        return '\n'.join(content)
    
    def _extract_section_text(self, section_name: str, start_index: int, strategy: str) -> str:
        """Extract section text according to strategy."""
        paragraphs = []
        word_count = 0
        word_limits = {
            'FULL': float('inf'),
            'FIRST_800_WORDS': 800,
            'FIRST_300_WORDS': 300
        }
        limit = word_limits.get(strategy, 300)
        
        collecting = False
        for i, para in enumerate(self.doc.Paragraphs):
            try:
                text = para.Range.Text.strip()
                
                if i == start_index:
                    collecting = True
                    continue
                
                if collecting:
                    if self._is_section_header(text):
                        break
                    
                    if text:
                        paragraphs.append(text)
                        word_count += len(text.split())
                        
                        if word_count >= limit:
                            paragraphs.append("\n[... resto omitido ...]")
                            break
            except:
                continue
        
        return '\n\n'.join(paragraphs)
           
    def generate_report_v06(self):
        """Generate comprehensive v0.6 report."""

        report = []
        SEP = "=" * 70

        # ---- local helpers (method-scoped) ----
        def header(title):
            report.append("\n" + SEP)
            report.append(title)
            report.append(SEP)

        def bullets(items, formatter=str):
            for item in items:
                report.append(f"  • {formatter(item)}")

        def severity_block(severity, items, formatter=str):
            if not items:
                return
            report.append(f"\n{severity.value}:")
            bullets(items, formatter)

        # === REPORT HEADER ===
        report.append(SEP)
        report.append("SILVINA v0.6 - REPORTE COMPLETO")
        report.append(SEP)
        report.append(f"Documento: {os.path.basename(self.filepath)}")
        report.append(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        report.append(f"Caracteres: {self.get_character_count():,}")

        # === ARTICLE CLASSIFICATION ===
        if self.article_classification:
            cls = self.article_classification

            header("CLASIFICACIÓN DE ARTÍCULO (Determinística)")
            report.append(f"Tipo: {cls['type'].value}")
            report.append(f"Confianza: {cls['confidence'].upper()}")
            report.append(f"Puntuación: {cls['score']}/10")

            severity_block(
                SeverityLevel.CRITICO,
                cls['reasons']['critical']
            )

            if cls['reasons']['positive']:
                report.append("\n✅ Indicadores Positivos:")
                bullets(cls['reasons']['positive'])

        # === IMRyD STRUCTURE ===
        if self.sections:
            header("ESTRUCTURA IMRyD")
            report.append(f"Secciones detectadas: {len(self.sections)}/5")

            for section in sorted(self.sections, key=lambda s: s.expected_order):
                report.append(
                    f"  {section.expected_order}. "
                    f"{section.name.title()} - {section.word_count} palabras"
                )

            validator = StructureValidator()
            issues = validator.validate(self.sections)

            severity_block(
                SeverityLevel.CRITICO,
                issues['missing_sections'],
                formatter=lambda s: s.title()
            )

            severity_block(
                SeverityLevel.CRITICO,
                issues['out_of_order'],
                formatter=lambda p: f"{p[0].title()} antes de {p[1].title()}"
            )

            severity_block(
                SeverityLevel.ADVERTENCIA,
                issues['too_short'],
                formatter=lambda s: (
                    f"{s['section'].title()}: "
                    f"{s['current']} palabras (mín: {s['minimum']})"
                )
            )

        # === CITATION INTEGRITY ===
        if self.citations or self.references:
            matcher = CitationMatcher(self.citations, self.references)
            report.append(matcher.generate_report(self.section_type))

        # === REFERENCE VALIDATION ===
        header("VALIDACIÓN DE REFERENCIAS APA")

        if self.references:
            validated = [(ref, ref.is_valid()) for ref in self.references]
            valid_count = sum(1 for _, ok in validated if ok)
            invalid = [(i, ref) for i, (ref, ok) in enumerate(validated, 1) if not ok]

            report.append(f"Total: {len(self.references)}")
            report.append(f"✅ Válidas: {valid_count}")
            report.append(f"❌ Con problemas: {len(invalid)}")

            if invalid:
                report.append("\nDetalle de Referencias con Problemas:")
                for i, ref in invalid:
                    rep = ref.get_validation_report()
                    report.append(f"\n{i}. {rep['text']}")

                    if not rep['valid_author']:
                        report.append("   ⚠️ Formato de autor incorrecto")
                    if not rep['valid_year']:
                        report.append("   ⚠️ Año no encontrado")
                    if not rep['valid_conjuncion']:
                        report.append(f"   ⚠️ {rep['error_conjuncion']}")
        else:
            report.append("⚠️  No se encontraron referencias para validar")

        report.append("\n" + SEP)
        return "\n".join(report)

                    
    def _add_section_header(self, doc, text):
        """Agrega un encabezado de sección con formato uniforme."""
        doc.add_paragraph()
        heading = doc.add_heading(text, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(10, 118, 184)
       
    def export_to_word(self, output_path: str):
        """Export comprehensive report to formatted Word document."""
        doc_export = DocxDocument()
    
        # ============================================================
        # DOCUMENT SETTINGS
        # ============================================================
        sections = doc_export.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
            # Header
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = "SILVINA – Asistente Editorial | Facultad Militar Conjunta – UNDEF"
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = header_para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(107, 113, 120)
            
            # Footer with page number
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            footer_para.text = "Página "
            run = footer_para.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(107, 113, 120)
            # Add page number field
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
    
            # Set default font
            style = doc_export.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
    
            # ============================================================
            # TITLE (Italic subtitle only)
            # ============================================================
            title = doc_export.add_paragraph('Reporte Editorial Completo')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.runs[0]
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(14)
            title_run.italic = True
            
            doc_export.add_paragraph()  # Spacer
            
            # ============================================================
            # DOCUMENT INFO TABLE
            # ============================================================
            info_table = doc_export.add_table(rows=4, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_table.rows[0].cells[0].text = 'Documento'
            info_table.rows[0].cells[1].text = os.path.basename(self.filepath)
            info_table.rows[1].cells[0].text = 'Fecha'
            info_table.rows[1].cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M')
            info_table.rows[2].cells[0].text = 'Caracteres'
            info_table.rows[2].cells[1].text = f"{self.get_character_count():,}"
            info_table.rows[3].cells[0].text = 'Palabras'
            info_table.rows[3].cells[1].text = f"{self.word_count:,}"
            
            # Format all cells
            for row in info_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
            
            doc_export.add_paragraph()  # Spacer
            
            # ============================================================
            # SECTION 1: CLASSIFICATION
            # ============================================================
            heading = doc_export.add_heading('CLASIFICACIÓN DE ARTÍCULO', level=1)
            heading.runs[0].font.name = 'Times New Roman'
            
            if self.article_classification:
                cls = self.article_classification
                
                # Type (bold inline)
                p = doc_export.add_paragraph()
                p.add_run('Tipo: ').bold = True
                p.add_run(cls['type'].value).bold = True
                
                # Confidence & Score
                p = doc_export.add_paragraph()
                p.add_run(f"Confianza: {cls['confidence'].upper()}")
                
                p = doc_export.add_paragraph()
                p.add_run(f"Puntuación: {cls['score']}/10")
                
                # CRITICAL ISSUES
                if cls['reasons']['critical']:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('🔴 CRÍTICO:')
                    run.bold = True
                    
                    for issue in cls['reasons']['critical']:
                        p = doc_export.add_paragraph(f"• {issue}")
                
                # POSITIVE INDICATORS
                if cls['reasons']['positive']:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('✅ Indicadores Positivos:')
                    run.bold = True
                    
                    for item in cls['reasons']['positive']:
                        p = doc_export.add_paragraph(f"• {item}")
            
            # ============================================================
            # SECTION 2: IMRYD STRUCTURE
            # ============================================================
            heading = doc_export.add_heading('ESTRUCTURA IMRyD', level=1)
            heading.runs[0].font.name = 'Times New Roman'
            
            p = doc_export.add_paragraph()
            p.add_run(f"Secciones detectadas: {len(self.sections)}/5")
            
            # List detected sections
            if self.sections:
                doc_export.add_paragraph()
                for section in sorted(self.sections, key=lambda s: s.expected_order):
                    p = doc_export.add_paragraph(
                        f"{section.expected_order}. {section.name.title()} - {section.word_count} palabras"
                    )
                
                # Validate structure
                validator = StructureValidator()
                issues = validator.validate(self.sections)
                
                # MISSING SECTIONS
                if issues['missing_sections']:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('🔴 CRÍTICO: Secciones Faltantes')
                    run.bold = True
                    
                    for missing in issues['missing_sections']:
                        p = doc_export.add_paragraph(f"• {missing.title()}")
                
                # OUT OF ORDER
                if issues['out_of_order']:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('🔴 CRÍTICO: Orden Incorrecto')
                    run.bold = True
                    
                    for sec1, sec2 in issues['out_of_order']:
                        p = doc_export.add_paragraph(f"• {sec1.title()} antes de {sec2.title()}")
                
                # TOO SHORT
                if issues['too_short']:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('🟡 ADVERTENCIA: Secciones Cortas')
                    run.bold = True
                    
                    for short in issues['too_short']:
                        p = doc_export.add_paragraph(
                            f"• {short['section'].title()}: {short['current']} palabras (mín: {short['minimum']})"
                        )
            else:
                p = doc_export.add_paragraph('⚠️ No se detectaron secciones IMRyD')
            
            # ============================================================
            # SECTION 3: CITATION INTEGRITY
            # ============================================================
            heading = doc_export.add_heading('INTEGRIDAD DE CITAS', level=1)
            heading.runs[0].font.name = 'Times New Roman'
            
            p = doc_export.add_paragraph()
            p.add_run(f"Citas en texto: {len(self.citations)}")
            
            p = doc_export.add_paragraph()
            p.add_run(f"Referencias bibliográficas: {len(self.references)}")
            
            p = doc_export.add_paragraph()
            p.add_run(f"Tipo de sección: {self.section_type.upper()}")
            
            if self.citations or self.references:
                matcher = CitationMatcher(self.citations, self.references)
                
                doc_export.add_paragraph()
                
                # ORPHANED REFERENCES
                orphaned_refs = matcher.find_orphaned_references()
                if orphaned_refs:
                    p = doc_export.add_paragraph()
                    
                    if self.section_type == "Referencias":
                        run = p.add_run('🟡 ADVERTENCIA: Referencias Sin Citar')
                    else:
                        run = p.add_run('🔵 INFORMATIVO: Referencias Sin Citar')
                    run.bold = True
                    
                    if self.section_type == "Referencias":
                        msg = "En 'Referencias', se espera citar todas las entradas."
                    else:
                        msg = "En 'Bibliografía', es aceptable incluir fuentes consultadas."
                    
                    p = doc_export.add_paragraph(msg)
                    
                    p = doc_export.add_paragraph(f"Encontradas {len(orphaned_refs)} referencias no citadas:")
                    
                    doc_export.add_paragraph()
                    for ref in orphaned_refs[:10]:
                        p = doc_export.add_paragraph(f"• {ref.text[:60]}...")
                
                # ORPHANED CITATIONS
                orphaned_cits = matcher.find_orphaned_citations()
                if orphaned_cits:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('🔴 CRÍTICO: Citas Sin Referencia')
                    run.bold = True
                    
                    p = doc_export.add_paragraph(f"Encontradas {len(orphaned_cits)} citas sin entrada bibliográfica:")
                    
                    doc_export.add_paragraph()
                    for cit in orphaned_cits[:10]:
                        p = doc_export.add_paragraph(f"• {cit}")
                
                # SUCCESS MESSAGE
                if not orphaned_cits and not orphaned_refs:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('✅ Sistema de citación íntegro')
                    run.bold = True
                elif not orphaned_cits:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('✅ Todas las citas tienen referencia válida')
                    run.bold = True
            
            # ============================================================
            # SECTION 4: REFERENCE VALIDATION
            # ============================================================
            heading = doc_export.add_heading('VALIDACIÓN DE REFERENCIAS APA', level=1)
            heading.runs[0].font.name = 'Times New Roman'
            
            if len(self.references) > 0:
                valid_count = sum(1 for ref in self.references if ref.is_valid())
                invalid_count = len(self.references) - valid_count
                
                p = doc_export.add_paragraph()
                p.add_run(f"Total: {len(self.references)}")
                
                p = doc_export.add_paragraph()
                p.add_run(f"✅ Válidas: {valid_count}")
                
                if invalid_count > 0:
                    p = doc_export.add_paragraph()
                    p.add_run(f"❌ Con problemas: {invalid_count}")
                else:
                    p = doc_export.add_paragraph()
                    p.add_run(f"❌ Con problemas: 0")
                
                # Show invalid references
                if invalid_count > 0:
                    doc_export.add_paragraph()
                    p = doc_export.add_paragraph()
                    run = p.add_run('Detalle de Referencias con Problemas:')
                    run.bold = True
                    
                    for i, ref in enumerate(self.references, 1):
                        if not ref.is_valid():
                            rep = ref.get_validation_report()
                            
                            doc_export.add_paragraph()
                            p = doc_export.add_paragraph(f"{i}. {rep['text']}")
                            
                            if not rep['valid_author']:
                                p = doc_export.add_paragraph("   ⚠️ Formato de autor incorrecto")
                            if not rep['valid_year']:
                                p = doc_export.add_paragraph("   ⚠️ Año no encontrado")
                            if not rep['valid_conjuncion']:
                                p = doc_export.add_paragraph(f"   ⚠️ {rep['error_conjuncion']}")
            else:
                p = doc_export.add_paragraph('⚠️ No se encontraron referencias para validar')
            
            # ============================================================
            # SECTION 5: LLM TIER 2 PLACEHOLDER
            # ============================================================
            doc_export.add_page_break()
            heading = doc_export.add_heading('TIER 2: ANÁLISIS DE CALIDAD (LLM)', level=1)
            heading.runs[0].font.name = 'Times New Roman'

            p = doc_export.add_paragraph('Análisis LLM pendiente.')
            # Set line spacing to 1.0
            p.paragraph_format.line_spacing = 1.0

            # Save
            doc_export.save(output_path)
            print(f"📄 Reporte Word guardado: {output_path}")
            
                        
    def close(self):
        """Close Word connection."""
        try:
            if self.doc:
                self.doc.Close(SaveChanges=False)
            if self.word:
                self.word.Quit()
        except:
            pass


# ============================================================
# MAIN EXECUTION
# ============================================================

if __name__ == "__main__":
    print("\n" + "="*70)
    print("SILVINA v0.6 - ASISTENTE EDITORIAL REDESIGNED")
    print("="*70 + "\n")
    
    # Load document
    filepath = r"C:\Users\usuario\Desktop\Escudo cuantico_AB_25092025.docx"
    
    doc = Document(filepath)
    doc.load()
    
    # TIER 1: Deterministic validation
    doc.classify_article()
    tier1_report = doc.generate_report_v06()
    print(tier1_report)
   
    # TIER 2: LLM Quality Analysis
    print("\n" + "="*70)
    print(f"TIER 2: ANÁLISIS DE CALIDAD (Ollama - {OLLAMA_MODEL})")
    print("="*70)
           
    llm_content = doc.get_llm_analysis_content()
    llm_report = ""

    if llm_content:
        article_type = doc.article_classification['type'].value
        llm_report = analyze_with_ollama(llm_content, article_type, tier1_report)
        
        # Remove "Recomendaciones" section
        if "Recomendaciones" in llm_report or "**Recomendaciones**" in llm_report:
            llm_report = llm_report.split("Recomendaciones")[0].strip()
            llm_report = llm_report.split("**Recomendaciones**")[0].strip()
        
        print(llm_report)
    else:
        llm_report = "⚠️ No se pudo extraer contenido para análisis LLM"
        print(llm_report)
    
    # Save text report
    combined_report = tier1_report + "\n\n" + "="*70 + "\n"
    combined_report += "TIER 2: ANÁLISIS DE CALIDAD (Ollama - llama3-gradient:8b)\n"
    combined_report += "="*70 + "\n" + llm_report
    
    report_filename = f"reporte_silvina_v06_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    with open(report_filename, 'w', encoding='utf-8') as f:
        f.write(combined_report)
    
    print(f"\n💾 Reporte completo guardado: {report_filename}")
    
    # Export to Word
    word_filename = f"reporte_silvina_v06_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.export_to_word(word_filename)
    
    # Add LLM analysis to Word document (only if valid)
    if llm_report and "No se pudo extraer" not in llm_report and "No conecta" not in llm_report:
        try:
            word_doc = DocxDocument(word_filename)
            
            # Find and replace "Análisis LLM pendiente."
            for i in range(len(word_doc.paragraphs) - 1, -1, -1):
                if "Análisis LLM pendiente" in word_doc.paragraphs[i].text:
                    word_doc.paragraphs[i].text = ""
                    break
            
            # Add formatted LLM analysis
            p = word_doc.add_paragraph()
            run = p.add_run('Análisis')
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.0
            
            # Add LLM content
            llm_para = word_doc.add_paragraph(llm_report)
            llm_para.paragraph_format.line_spacing = 1.0
            for run in llm_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            word_doc.save(word_filename)
            print(f"✅ Análisis LLM agregado al reporte Word")
        except Exception as e:
            print(f"⚠️ Error agregando LLM al Word: {e}")
    
    doc.close()



